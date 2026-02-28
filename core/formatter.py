# core/formatter.py
import re
from typing import Dict, List, Set
from collections import Counter

from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .parser import Block
from .spec import Spec
from .docx_utils import (
    copy_run_style,
    delete_paragraph,
    is_effectively_blank_paragraph,
    iter_all_paragraphs,
    normalize_mixed_runs,
    set_run_fonts,
)


# =========================
# Regex rules
# =========================

RE_SUBTITLE_CN = re.compile(r"^\s*（[一二三四五六七八九十]+）")  # （一）
RE_CAPTION = re.compile(
    r"^\s*(图|表|Figure|Fig\.|Table)\s*[\d一二三四五六七八九十]+([\-–—]\d+)?",
    re.IGNORECASE
)

RE_CN_ENUM = re.compile(r"^\s*[一二三四五六七八九十百千万]+、")  # 接受所有中文数字字符前缀（含百千万），不强制校验组合合法性
RE_NUM_DOT = re.compile(r"^\s*\d+(\.\d+){0,3}\s+")
RE_ABSTRACT = re.compile(r"^\s*(摘要|abstract)\s*[:：]?\s*", re.IGNORECASE)
RE_KEYWORD = re.compile(r"^\s*(关键词|关键字|keywords?)\s*[:：]?\s*", re.IGNORECASE)
RE_REFERENCE = re.compile(r"^\s*(参考文献|references?|bibliography)\s*$", re.IGNORECASE)
ROLE_LABELS_FALLBACK_TO_RULE = {"blank", "unknown"}

# 段内多行结构（避免误判标题）
RE_MULTILINE_NUM = re.compile(r"\n\s*\d+(\.\d+)*\s+")
RE_MULTILINE_SUB = re.compile(r"\n\s*（[一二三四五六七八九十]+）")


# =========================
# Helpers
# =========================

def is_list_paragraph(p: Paragraph) -> bool:
    """True if paragraph is a Word numbered/bulleted list (numPr)."""
    try:
        ppr = p._p.pPr
        return bool(ppr is not None and getattr(ppr, 'numPr', None) is not None)
    except Exception:
        return False

def looks_like_multiline_numbered_block(text: str) -> bool:
    t = text or ""
    return bool(RE_MULTILINE_NUM.search(t)) or bool(RE_MULTILINE_SUB.search(t))


def _clear_paragraph_runs(p):
    """彻底清空 run（含 br），用于重建文本。"""
    for child in list(p._p):
        if child.tag.endswith("}r"):
            p._p.remove(child)


def _strip_trailing_newlines_in_paragraph(p):
    txt = p.text or ""
    new_txt = txt.rstrip("\r\n")
    if new_txt == txt:
        return
    _clear_paragraph_runs(p)
    p.add_run(new_txt)


def _insert_paragraph_after(p, text: str):
    """在段落 p 后插入新段落（稳定版），并写入 text。"""
    new_p = OxmlElement('w:p')
    p._p.addnext(new_p)
    new_para = Paragraph(new_p, p._parent)
    new_para.add_run(text)
    return new_para


# =========================
# Role detection
# =========================

def detect_role(paragraph) -> str:
    """
    Detect the structural role of a paragraph.

    Returns one of:
      blank / h1 / h2 / h3 / caption / abstract / keyword / reference /
      list_item / footer / body

    Rules applied in priority order:
      1. blank (empty / whitespace-only)
      2. body  (multi-line numbered block – avoids misclassifying as heading)
      3. Word heading/footer styles
      4. Abstract / keyword / reference text patterns
      5. Caption text pattern (before list_item so captioned lists stay as caption)
      6. Word list paragraph (numPr)
      7. （一）-style sub-heading → h3
      8. 第X章/节/条 patterns → h1/h2/h3
      9. Chinese numeral enum (一、二、…) → h2
      10. Numeric outline (1. / 1.1) → h2/h3
      11. Fallback → body
    """
    if is_effectively_blank_paragraph(paragraph):
        return "blank"

    text = paragraph.text or ""

    # 段内多行编号块强制当正文，避免误判标题
    if looks_like_multiline_numbered_block(text):
        return "body"

    # 优先尊重 Word 标题样式
    style_name = ""
    try:
        style_name = (paragraph.style.name or "").lower()
    except Exception:
        style_name = ""
    if "heading 1" in style_name or "标题 1" in style_name:
        return "h1"
    if "heading 2" in style_name or "标题 2" in style_name:
        return "h2"
    if "heading 3" in style_name or "标题 3" in style_name:
        return "h3"
    if "footer" in style_name or "页脚" in style_name:
        return "footer"

    t = text.strip()

    if RE_ABSTRACT.match(t):
        return "abstract"
    if RE_KEYWORD.match(t):
        return "keyword"
    if RE_REFERENCE.match(t):
        return "reference"
    if RE_CAPTION.match(t):
        return "caption"
    if is_list_paragraph(paragraph):
        return "list_item"

    if RE_SUBTITLE_CN.match(t):
        return "h3"

    if t.startswith("第") and "章" in t[:12]:
        return "h1"
    if t.startswith("第") and "节" in t[:12]:
        return "h2"
    if t.startswith("第") and "条" in t[:12]:
        return "h3"
    if RE_CN_ENUM.match(t):
        return "h2"
    if RE_NUM_DOT.match(t):
        depth = t.split()[0].count(".")
        return "h2" if depth <= 0 else "h3"

    return "body"


# =========================
# Formatting helpers
# =========================

def _apply_paragraph_common(p, line_spacing: float, space_before_pt: float, space_after_pt: float):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def _resolve_alignment(name: str):
    """将配置里的对齐字符串映射为 python-docx 枚举。"""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return mapping.get((name or "").strip().lower())


def _apply_runs_font(p, zh_font: str, en_font: str, size_pt: float, force_bold=None):
    # 先把中英混合 run 拆开，再逐 run 写入完整字体映射
    normalize_mixed_runs(p)
    for run in p.runs:
        run.font.size = Pt(size_pt)
        if force_bold is not None:
            run.font.bold = bool(force_bold)
        set_run_fonts(run, zh_font=zh_font, en_font=en_font)


def _first_line_indent_pt(chars: int, font_size_pt: float) -> Pt:
    # 近似：1 个中文字符宽度 ≈ 1 个字号(pt)
    return Pt(chars * font_size_pt)


def _cleanup_consecutive_blanks(doc, max_keep: int) -> int:
    """压缩连续空段：最多保留 max_keep 个（0=全删）。返回删除的空段数量。"""
    blank_run = 0
    to_delete = []
    last_parent = None
    for p in list(iter_all_paragraphs(doc)):
        cur_parent = p._element.getparent()
        if cur_parent is not last_parent:
            # 不跨容器（正文/单元格）累计空段，避免误删
            blank_run = 0
            last_parent = cur_parent
        if is_effectively_blank_paragraph(p):
            blank_run += 1
            if blank_run > max_keep:
                to_delete.append(p)
        else:
            blank_run = 0

    deleted = 0
    # 倒序删除，避免重排带来的错删
    for p in reversed(to_delete):
        delete_paragraph(p)
        deleted += 1
    return deleted


def _delete_blanks_after_roles(doc, roles: Set[str], role_getter=None) -> int:
    """删除“标题/题注后紧跟的所有空段”。返回删除数量。"""
    if role_getter is None:
        role_getter = detect_role

    to_delete = []
    paras = list(iter_all_paragraphs(doc))
    i = 0
    while i < len(paras):
        cur = paras[i]
        cur_role = role_getter(cur)
        if cur_role in roles:
            cur_parent = cur._element.getparent()
            j = i + 1
            while j < len(paras):
                nxt = paras[j]
                if nxt._element.getparent() is not cur_parent:
                    break
                if not is_effectively_blank_paragraph(nxt):
                    break
                to_delete.append(nxt)
                j += 1
        i += 1

    for p in reversed(to_delete):
        delete_paragraph(p)
    return len(to_delete)


def _split_body_paragraphs_on_linebreaks(doc, role_getter=None, on_new_paragraph=None) -> int:
    """
    关键修复：
    把正文段落里的 '\n'（通常是 Shift+Enter 软回车）拆成多个段落，
    这样每一条（比如 \n1. \n2.）都能获得“首行缩进”。

    - role_getter: 用于判断某段是否为 body（默认 detect_role）。
    - on_new_paragraph(parent, child): 可选回调；当从 parent 拆出 child 时调用。
      用于“继承标签/元数据”等（例如：child 的 role 继承 parent）。

    返回：新增段落数量（插入了多少个新段落）。
    """
    if role_getter is None:
        role_getter = detect_role

    created = 0
    # 使用快照：覆盖正文与表格段落，同时避免边插入边遍历造成错位
    for p in list(iter_all_paragraphs(doc)):
        if is_effectively_blank_paragraph(p):
            continue

        # 只拆正文；标题/题注不拆，避免破坏结构
        role = role_getter(p)
        if role != "body":
            continue

        text = p.text or ""
        if "\n" not in text:
            continue

        # 保留每一行对应的“源 run 样式”（颜色/粗斜体），避免拆段后样式丢失
        raw_runs = list(p.runs)
        line_parts = [[]]
        for src_run in raw_runs:
            parts = (src_run.text or "").split("\n")
            for idx, part in enumerate(parts):
                if part:
                    line_parts[-1].append((part, src_run))
                if idx < len(parts) - 1:
                    line_parts.append([])

        resolved_lines = []
        for parts in line_parts:
            line_text = "".join(seg for seg, _ in parts).strip()
            if not line_text:
                continue
            style_run = parts[0][1] if parts else (raw_runs[0] if raw_runs else None)
            resolved_lines.append((line_text, style_run))

        if len(resolved_lines) <= 1:
            continue

        # 当前段落替换成第一行并继承该行样式
        _clear_paragraph_runs(p)
        first_text, first_style_run = resolved_lines[0]
        first_run = p.add_run(first_text)
        if first_style_run is not None:
            copy_run_style(first_style_run, first_run)

        # 后续行插入为新段落，复制段落样式与该行首 run 样式
        prev = p
        for ln, style_run in resolved_lines[1:]:
            new_p = _insert_paragraph_after(prev, ln)
            created += 1
            try:
                new_p.style = p.style
            except Exception:
                pass
            if style_run is not None and new_p.runs:
                copy_run_style(style_run, new_p.runs[0])

            if on_new_paragraph is not None:
                try:
                    on_new_paragraph(p, new_p)
                except Exception:
                    # 回调失败不应影响主流程
                    pass

            prev = new_p
    return created


def apply_formatting(doc, blocks: List[Block], labels: Dict[int, str], spec: Spec):
    """
    MVP：不处理真编号/列表结构，只做视觉排版（首行缩进、字体、字号、行距、段距、空行清理）。
    且会把正文中的 '\\n' 拆成多个段落，保证每条都缩进。

    返回：report(dict) —— 可直接 dump 为 JSON，用于“诊断/修复报告”。
    """
    cfg = spec.raw
    zh_font = cfg["fonts"]["zh"]
    en_font = cfg["fonts"]["en"]

    body_cfg = cfg["body"]
    body_size = float(body_cfg["font_size_pt"])
    body_line_spacing = float(body_cfg["line_spacing"])
    body_before = float(body_cfg["space_before_pt"])
    body_after = float(body_cfg["space_after_pt"])
    first_line_chars = int(body_cfg["first_line_chars"])

    heading_cfg = cfg["heading"]
    caption_cfg = cfg.get("caption", None)
    paragraph_cfg = cfg.get("paragraph", {})

    cleanup_cfg = cfg.get("cleanup", {})
    max_blank_keep = int(cleanup_cfg.get("max_consecutive_blank_paragraphs", 1))
    remove_blank_after_roles = set(cleanup_cfg.get("remove_blank_after_roles", ["h1", "h2", "h3", "caption"]))

    list_cfg = cfg.get("list_item", {})
    list_left_indent = float(list_cfg.get("left_indent_pt", 18))
    list_hanging_indent = float(list_cfg.get("hanging_indent_pt", 18))

    body_alignment = _resolve_alignment(paragraph_cfg.get("alignment", "justify"))

    # ====== 角色映射：优先 labels，缺失才 fallback detect_role ======
    # 关键点：用底层 CT_P XML 元素做 key（而不是 Paragraph 包装对象），
    # 因为每次调用 iter_all_paragraphs 都会创建新的 Paragraph 包装对象，
    # 若用对象本身做 key 会导致 label_by_elem 查找永远失败。
    orig_paras = iter_all_paragraphs(doc)
    para_by_index = {i: p for i, p in enumerate(orig_paras)}
    label_by_elem: Dict = {}  # CT_P element -> role str

    for b in blocks:
        role = labels.get(b.block_id)
        if not role or role in ROLE_LABELS_FALLBACK_TO_RULE:
            continue
        p = para_by_index.get(b.paragraph_index)
        if p is not None:
            label_by_elem[p._p] = role

    def get_role(p: Paragraph) -> str:
        return label_by_elem.get(p._p) or detect_role(p)

    # ====== Report（诊断/动作统计/可解释输出）======
    label_source = labels.get("_source", "unknown")
    # 仅统计 blocks 对应的标签，忽略 labels 里的元信息键
    label_list = [labels.get(b.block_id, "unknown") for b in blocks]
    label_counts = dict(Counter(label_list))

    # labels 覆盖率（仅针对 blocks；unknown/blank 视为未标注）
    total_blocks = len(blocks)
    labeled_blocks = 0
    for b in blocks:
        lab = labels.get(b.block_id)
        if lab and lab not in ("blank", "unknown"):
            labeled_blocks += 1
    coverage_rate = (labeled_blocks / total_blocks) if total_blocks else 0.0

    # label vs fallback 规则的一致性（仅针对原始段落）
    mismatch = 0
    compared = 0
    mismatch_examples = []
    for b in blocks:
        p = para_by_index.get(b.paragraph_index)
        if p is None:
            continue
        lab = labels.get(b.block_id)
        if not lab or lab == "blank":
            continue
        compared += 1
        det = detect_role(p)
        if det != lab:
            mismatch += 1
            if len(mismatch_examples) < 5:
                mismatch_examples.append({
                    "paragraph_index": b.paragraph_index,
                    "block_id": b.block_id,
                    "label": lab,
                    "detect_role": det,
                    "text": (p.text or "")[:120],
                })

    list_para_count = sum(1 for p in orig_paras if is_list_paragraph(p))

    report = {
        "meta": {
            "paragraphs_before": len(orig_paras),
            "blank_paragraphs_before": sum(1 for p in orig_paras if is_effectively_blank_paragraph(p)),
            "list_paragraphs_before": list_para_count,
        },
        "labels": {
            "source": label_source,
            "counts": label_counts,
            "coverage": {
                "total_blocks": total_blocks,
                "labeled_blocks": labeled_blocks,
                "coverage_rate": coverage_rate,
            },
            "consistency": {
                "compared": compared,
                "mismatched": mismatch,
                "mismatch_rate": (mismatch / compared) if compared else 0.0,
                "mismatch_examples": mismatch_examples,
            },
        },
        "plan_executed": [
            "cleanup_consecutive_blanks",
            "delete_blanks_after_titles",
            "split_body_paragraphs_on_linebreaks",
            "apply_formatting",
        ],
        "actions": {},
        "formatted": {"counts": {}},
        "warnings": [],
    }
    # 1) 空段压缩/清理
    deleted_consecutive = _cleanup_consecutive_blanks(doc, max_blank_keep)
    report["actions"]["cleanup_consecutive_blanks_deleted"] = deleted_consecutive
    report["actions"]["cleanup_consecutive_blank_keep"] = max_blank_keep

    # 2) 标题/题注后空段删光
    deleted_after_roles = _delete_blanks_after_roles(doc, roles=remove_blank_after_roles, role_getter=get_role)
    report["actions"]["delete_blanks_after_titles_deleted"] = deleted_after_roles

    # 3) 核心修复：拆正文段落里的软回车换行（\n）
    # 先做一次“预估/统计”，便于 report 诊断
    split_affected = 0
    split_max_lines = 0
    split_estimated_new = 0
    for p in orig_paras:
        if is_effectively_blank_paragraph(p):
            continue
        if get_role(p) != "body":
            continue
        t = p.text or ""
        if "\n" not in t:
            continue
        lines = [ln.strip() for ln in t.split("\n") if ln.strip()]
        if len(lines) <= 1:
            continue
        split_affected += 1
        split_max_lines = max(split_max_lines, len(lines))
        split_estimated_new += (len(lines) - 1)

    new_paras_from_split: List[Paragraph] = []
    def _inherit_label(parent_p, child_p):
        # 让拆分出来的新段落继承原段落的标签（避免 fallback 造成标签不一致）
        if parent_p._p in label_by_elem and child_p._p not in label_by_elem:
            label_by_elem[child_p._p] = label_by_elem[parent_p._p]
        new_paras_from_split.append(child_p)

    created_by_split = _split_body_paragraphs_on_linebreaks(doc, role_getter=get_role, on_new_paragraph=_inherit_label)
    report["actions"]["split_body_new_paragraphs_created"] = created_by_split
    report["actions"]["split_body_original_paragraphs_affected"] = split_affected
    report["actions"]["split_body_max_lines_in_one_paragraph"] = split_max_lines
    report["actions"]["split_body_estimated_new_paragraphs"] = split_estimated_new

    # 4) 套格式
    formatted_counter = Counter()
    for p in iter_all_paragraphs(doc):
        if is_effectively_blank_paragraph(p):
            continue

        role = get_role(p)

        # 标题/题注：去掉段尾多余换行
        if role in ("h1", "h2", "h3", "caption"):
            _strip_trailing_newlines_in_paragraph(p)

        if role == "body":
            _apply_paragraph_common(p, body_line_spacing, body_before, body_after)

            # 普通正文与 Word 列表正文分别处理缩进，避免相互覆盖
            if is_list_paragraph(p):
                p.paragraph_format.left_indent = Pt(list_left_indent)
                p.paragraph_format.hanging_indent = Pt(list_hanging_indent)
                p.paragraph_format.first_line_indent = Pt(0)
                formatted_counter["list_body"] += 1
            else:
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.hanging_indent = Pt(0)
                p.paragraph_format.first_line_indent = _first_line_indent_pt(first_line_chars, body_size)
                formatted_counter["body"] += 1

            if body_alignment is not None:
                p.paragraph_format.alignment = body_alignment
            _apply_runs_font(p, zh_font, en_font, size_pt=body_size, force_bold=None)

        elif role in ("h1", "h2", "h3"):
            hc = heading_cfg[role]
            size = float(hc["font_size_pt"])
            bold = bool(hc["bold"])
            before = float(hc["space_before_pt"])
            after = float(hc["space_after_pt"])
            heading_align = _resolve_alignment(hc.get("alignment", "left"))

            _apply_paragraph_common(p, body_line_spacing, before, after)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.hanging_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            if heading_align is not None:
                p.paragraph_format.alignment = heading_align
            _apply_runs_font(p, zh_font, en_font, size_pt=size, force_bold=bold)
            formatted_counter[role] += 1

        elif role == "caption":
            if caption_cfg:
                size = float(caption_cfg.get("font_size_pt", body_size))
                bold = bool(caption_cfg.get("bold", False))
                before = float(caption_cfg.get("space_before_pt", body_before))
                after = float(caption_cfg.get("space_after_pt", body_after))

                _apply_paragraph_common(p, body_line_spacing, before, after)
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.hanging_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
                cap_align = _resolve_alignment(caption_cfg.get("alignment", "center" if caption_cfg.get("center", True) else "left"))
                if cap_align is not None:
                    p.paragraph_format.alignment = cap_align
                _apply_runs_font(p, zh_font, en_font, size_pt=size, force_bold=bold)
            else:
                _apply_paragraph_common(p, body_line_spacing, body_before, body_after)
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.hanging_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
                if body_alignment is not None:
                    p.paragraph_format.alignment = body_alignment
                _apply_runs_font(p, zh_font, en_font, size_pt=body_size, force_bold=None)

            formatted_counter["caption"] += 1

        elif role in ("abstract", "keyword", "reference", "footer", "list_item"):
            rc = cfg.get(role, {})
            size = float(rc.get("font_size_pt", body_size))
            bold = bool(rc.get("bold", False))
            italic = bool(rc.get("italic", False))
            before = float(rc.get("space_before_pt", body_before))
            after = float(rc.get("space_after_pt", body_after))
            flc = int(rc.get("first_line_chars", 0))
            hanging = float(rc.get("hanging_indent_pt", 0))
            role_align = _resolve_alignment(rc.get("alignment", "justify"))

            _apply_paragraph_common(p, body_line_spacing, before, after)
            if hanging:
                # 悬挂缩进：所有行左缩进 hanging pt，首行回缩 -hanging pt（首行从页边起）
                p.paragraph_format.left_indent = Pt(hanging)
                p.paragraph_format.first_line_indent = Pt(-hanging)
            elif flc:
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.first_line_indent = _first_line_indent_pt(flc, size)
            else:
                p.paragraph_format.left_indent = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
            if role_align is not None:
                p.paragraph_format.alignment = role_align

            normalize_mixed_runs(p)
            for run in p.runs:
                run.font.size = Pt(size)
                run.font.bold = bold
                run.font.italic = italic
                set_run_fonts(run, zh_font=zh_font, en_font=en_font)

            formatted_counter[role] += 1

        else:
            # unknown：当正文处理，尽量不让段落漏掉缩进/字体统一
            _apply_paragraph_common(p, body_line_spacing, body_before, body_after)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.hanging_indent = Pt(0)
            p.paragraph_format.first_line_indent = _first_line_indent_pt(first_line_chars, body_size)
            if body_alignment is not None:
                p.paragraph_format.alignment = body_alignment
            _apply_runs_font(p, zh_font, en_font, size_pt=body_size, force_bold=None)
            formatted_counter["unknown_as_body"] += 1


    # ====== 追加诊断提示（让报告更“智能体”）======
    # 1) 软回车拆段提示
    if created_by_split > 0:
        report["warnings"].append(
            f"检测到正文段落含软回车(\\n)，已自动拆分为独立段落：新增 {created_by_split} 段。"
        )

    # 2) 列表段落提示（Word 真编号/项目符号可能覆盖缩进）
    if list_para_count > 0:
        report["warnings"].append(
            f"检测到 {list_para_count} 个 Word 列表/编号段落(numPr)。其缩进可能由列表级别控制，首行缩进不一致时建议单独处理列表缩进。"
        )

    # 3) 标签一致性提示
    if compared > 0 and mismatch > 0:
        rate = report["labels"]["consistency"]["mismatch_rate"]
        report["warnings"].append(
            f"标签与回退规则(detect_role)存在不一致：{mismatch}/{compared} (≈{rate:.0%})。建议统一以 labels 为准，并逐步收敛规则/LLM 标注。"
        )

    # 4) 标题层级提示：h3 前无 h2
    roles_seq = [get_role(p) for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    h2_seen = False
    orphan_h3 = 0
    for r in roles_seq:
        if r == "h2":
            h2_seen = True
        elif r == "h1":
            h2_seen = False
        elif r == "h3" and not h2_seen:
            orphan_h3 += 1
    if orphan_h3 > 0:
        report["warnings"].append(
            f"检测到 {orphan_h3} 个 h3 在其前方未出现 h2（可能层级断裂或标签误判）。"
        )

    # 额外统计：拆分新增段落的 role 分布（便于验证“新增段大多为正文”这一预期）
    try:
        new_role_counts = Counter(get_role(p) for p in new_paras_from_split if p is not None)
        report["actions"]["split_body_new_paragraph_roles"] = dict(new_role_counts)
    except Exception:
        report["actions"]["split_body_new_paragraph_roles"] = {}

    all_after = iter_all_paragraphs(doc)
    report["meta"]["paragraphs_after"] = len(all_after)
    report["meta"]["blank_paragraphs_after"] = sum(1 for p in all_after if is_effectively_blank_paragraph(p))
    report["formatted"]["counts"] = dict(formatted_counter)
    return report
