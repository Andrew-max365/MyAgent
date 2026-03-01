from docx import Document

from agent.doc_analyzer import DocAnalyzer


def test_doc_analyzer_extract_paragraphs_includes_table_paragraphs_in_flow_order():
    doc = Document()
    doc.add_paragraph("文档首段")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].text = "表格段落"
    doc.add_paragraph("文档尾段")

    paragraphs = DocAnalyzer.extract_paragraphs(doc)

    assert paragraphs == ["文档首段", "表格段落", "文档尾段"]
