# tests/test_detect_role_improvements.py
"""
Tests for improved detect_role patterns and heading alignment support.
"""
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.formatter import apply_formatting, detect_role
from core.spec import load_spec

SPECS_DIR = Path(__file__).resolve().parents[1] / "specs"


def _mark_paragraph_as_numbered_list(paragraph) -> None:
    """Mark a paragraph as a numbered list item via numPr XML."""
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = 1


def test_detect_role_di_tiao_is_h3():
    """「第X条」(法律条款) without numPr should be detected as h3."""
    doc = Document()
    assert detect_role(doc.add_paragraph("第一条 总则")) == "h3"
    assert detect_role(doc.add_paragraph("第十条 违约责任")) == "h3"
    assert detect_role(doc.add_paragraph("第百条 附则")) == "h3"


def test_detect_role_di_tiao_with_numpr_is_list_item():
    """「第X条」with Word numPr should still be list_item (numPr takes precedence)."""
    doc = Document()
    p = doc.add_paragraph("第一条 总则")
    _mark_paragraph_as_numbered_list(p)
    assert detect_role(p) == "list_item"


def test_detect_role_caption_before_list_item():
    """Caption-pattern text in a Word list paragraph should be classified as caption."""
    doc = Document()
    p = doc.add_paragraph("图1 系统架构图")
    _mark_paragraph_as_numbered_list(p)
    assert detect_role(p) == "caption"


def test_detect_role_cn_enum_extended_numerals():
    """中文序号如「百一、」「千一、」也应识别为 h2。"""
    doc = Document()
    assert detect_role(doc.add_paragraph("百一、总则")) == "h2"
    assert detect_role(doc.add_paragraph("千二、附则")) == "h2"


def test_heading_alignment_applied_in_formatting():
    """h1 should be centered, h2/h3 should be left-aligned after formatting."""
    from core.docx_utils import iter_all_paragraphs
    from core.parser import Block

    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    doc.add_paragraph("文章标题")    # h1
    doc.add_paragraph("第一节 引言")  # h2
    doc.add_paragraph("1.1.1 背景")   # h3

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {1: "h1", 2: "h2", 3: "h3", "_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    after = iter_all_paragraphs(doc)
    assert after[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert after[1].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert after[2].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_spec_heading_alignment_defaults():
    """Spec loader should fill heading alignment defaults: h1→center, h2/h3→left."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    assert spec.raw["heading"]["h1"]["alignment"] == "center"
    assert spec.raw["heading"]["h2"]["alignment"] == "left"
    assert spec.raw["heading"]["h3"]["alignment"] == "left"



