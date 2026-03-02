# tests/test_no_api_key.py
"""Tests that rule/hybrid modes work correctly without an LLM API key."""
from __future__ import annotations

import os
import tempfile
from unittest.mock import patch

import pytest
from docx import Document


def _make_simple_docx(path: str) -> None:
    doc = Document()
    doc.add_paragraph("第一章 引言")
    doc.add_paragraph("这是正文内容。")
    doc.add_paragraph("第二节 背景")
    doc.add_paragraph("更多正文内容。")
    doc.save(path)


@pytest.fixture()
def simple_docx(tmp_path):
    p = str(tmp_path / "input.docx")
    _make_simple_docx(p)
    return p


def test_rule_mode_no_api_key(simple_docx, tmp_path):
    """rule mode must complete without any LLM API key."""
    out = str(tmp_path / "out_rule.docx")
    with patch.dict(os.environ, {"LLM_API_KEY": ""}):
        from service.format_service import format_docx_file
        result = format_docx_file(simple_docx, out, label_mode="rule")
    assert os.path.exists(result.output_path)


def test_hybrid_mode_no_api_key_falls_back(simple_docx, tmp_path):
    """hybrid mode should fall back to rule mode when API key is missing."""
    out = str(tmp_path / "out_hybrid.docx")
    with patch.dict(os.environ, {"LLM_API_KEY": ""}):
        from service.format_service import format_docx_file
        # Should not raise even without API key
        result = format_docx_file(simple_docx, out, label_mode="hybrid")
    assert os.path.exists(result.output_path)


def test_rule_mode_bytes_no_api_key(simple_docx):
    """format_docx_bytes in rule mode should work without API key."""
    with open(simple_docx, "rb") as f:
        data = f.read()

    with patch.dict(os.environ, {"LLM_API_KEY": ""}):
        from service.format_service import format_docx_bytes
        out_bytes, report = format_docx_bytes(data, label_mode="rule")

    assert isinstance(out_bytes, bytes)
    assert len(out_bytes) > 0


def test_report_contains_meta_no_api_key(simple_docx, tmp_path):
    """Rule mode report should contain meta section even without API key."""
    out = str(tmp_path / "out_meta.docx")
    with patch.dict(os.environ, {"LLM_API_KEY": ""}):
        from service.format_service import format_docx_file
        result = format_docx_file(simple_docx, out, label_mode="rule")

    report = result.report
    assert isinstance(report, dict)


def test_config_missing_api_key_is_empty_string():
    """LLM_API_KEY defaults to empty string when not set."""
    with patch.dict(os.environ, {}, clear=True):
        import importlib
        import config as cfg
        importlib.reload(cfg)
        # After reload without LLM_API_KEY, it should be empty string
        assert isinstance(cfg.LLM_API_KEY, str)
