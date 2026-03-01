# tests/test_audit_fixes.py
"""
Tests for bugs found in the code audit:
  1. core/judge.py fallback RE_CN_ENUM now includes 百千万 (consistent with formatter.py)
  2. core/formatter.py unknown branch resets hanging_indent to Pt(0)
  3. hyperlink runs (URLs) get their fonts applied via iter_paragraph_runs
"""
from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.judge import rule_based_labels
from core.formatter import apply_formatting, detect_role
from core.spec import load_spec
from core.parser import Block
from core.docx_utils import iter_all_paragraphs

SPECS_DIR = Path(__file__).resolve().parents[1] / "specs"


# ---------------------------------------------------------------------------
# Fix 1: judge.py fallback RE_CN_ENUM includes 百千万
# ---------------------------------------------------------------------------

def test_judge_fallback_re_cn_enum_includes_bai_qian_wan():
    """Fallback rule_based_labels (no doc) must classify 百X/千X/万X enumerated
    headings as h2, consistent with formatter.py's detect_role."""
    from core.parser import Block

    blocks = [
        Block(block_id=1, kind="paragraph", text="百一、概论", paragraph_index=0),
        Block(block_id=2, kind="paragraph", text="千二、附则", paragraph_index=1),
        Block(block_id=3, kind="paragraph", text="万三、特别条款", paragraph_index=2),
    ]
    labels = rule_based_labels(blocks, doc=None)

    for b in blocks:
        assert labels[b.block_id] == "h2", (
            f"{b.text!r} should be 'h2' in fallback path, got {labels[b.block_id]!r}"
        )


# ---------------------------------------------------------------------------
# Fix 2 (now 2): apply_formatting unknown branch resets hanging_indent to Pt(0)
# ---------------------------------------------------------------------------

def _make_doc_blocks_labels(role_texts):
    """Helper: build (doc, blocks, labels) with one paragraph per (role, text)."""
    doc = Document()
    for _, text in role_texts:
        doc.add_paragraph(text)
    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {b.block_id: role for b, (role, _) in zip(blocks, role_texts)}
    labels["_source"] = "test"
    return doc, blocks, labels


def test_unknown_role_resets_hanging_indent():
    """A paragraph labelled 'unknown' must have left_indent reset to 0
    and first_line_indent set to a non-negative value (body-style indent)."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([("unknown", "这是未知角色的段落。")])

    # Pre-set a real hanging indent (left_indent > 0, first_line_indent < 0)
    para = iter_all_paragraphs(doc)[0]
    para.paragraph_format.left_indent = Pt(36)
    para.paragraph_format.first_line_indent = Pt(-36)

    apply_formatting(doc, blocks, labels, spec)

    after_para = iter_all_paragraphs(doc)[0]
    pf = after_para.paragraph_format

    assert pf.left_indent == Pt(0), (
        f"unknown branch should reset left_indent to 0, got {pf.left_indent}"
    )
    # first_line_indent should be positive (body first-line indent), not negative
    if pf.first_line_indent is not None:
        assert pf.first_line_indent >= 0, (
            f"unknown branch should not leave negative first_line_indent, got {pf.first_line_indent}"
        )


def test_unknown_role_hanging_indent_consistent_with_body():
    """Formatting a paragraph as 'unknown' should yield same left_indent and
    first_line_indent as 'body'."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    # unknown paragraph
    doc_u, blocks_u, labels_u = _make_doc_blocks_labels([("unknown", "未知段落内容。")])
    apply_formatting(doc_u, blocks_u, labels_u, spec)
    u_pf = iter_all_paragraphs(doc_u)[0].paragraph_format

    # body paragraph
    doc_b, blocks_b, labels_b = _make_doc_blocks_labels([("body", "正文段落内容。")])
    apply_formatting(doc_b, blocks_b, labels_b, spec)
    b_pf = iter_all_paragraphs(doc_b)[0].paragraph_format

    assert u_pf.left_indent == b_pf.left_indent, (
        f"unknown left_indent {u_pf.left_indent} != body {b_pf.left_indent}"
    )
    assert u_pf.first_line_indent == b_pf.first_line_indent, (
        f"unknown first_line_indent {u_pf.first_line_indent} != body {b_pf.first_line_indent}"
    )


# ---------------------------------------------------------------------------
# Fix 4 (now 4): hyperlink runs (URLs) get their fonts applied
# ---------------------------------------------------------------------------

def test_hyperlink_run_font_applied():
    """Runs inside w:hyperlink elements (e.g. URLs) must also get their font
    updated by apply_formatting — they were previously skipped because
    paragraph.runs does not include hyperlink-child runs."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from core.docx_utils import iter_paragraph_runs

    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    en_font = spec.raw["fonts"]["en"]

    doc = Document()
    p = doc.add_paragraph("正文前缀 ")

    # Manually embed a hyperlink run simulating URL text inside w:hyperlink
    hyperlink = OxmlElement("w:hyperlink")
    r_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = "https://example.com"
    r_elem.append(t_elem)
    hyperlink.append(r_elem)
    p._p.append(hyperlink)

    paras = iter_all_paragraphs(doc)
    blocks = [Block(block_id=1, kind="paragraph", text=paras[0].text, paragraph_index=0)]
    labels = {1: "body", "_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    # After formatting, ALL runs (including hyperlink runs) should have en_font
    all_runs = list(iter_paragraph_runs(iter_all_paragraphs(doc)[0]))
    assert len(all_runs) >= 1, "Expected at least one run after formatting"
    for run in all_runs:
        rpr = run._element.rPr
        assert rpr is not None, f"Run {run.text!r} has no rPr after formatting"
        rFonts = rpr.rFonts
        assert rFonts is not None, f"Run {run.text!r} has no rFonts after formatting"
        ascii_font = rFonts.get(qn("w:ascii"))
        assert ascii_font == en_font, (
            f"Run {run.text!r}: expected w:ascii={en_font!r}, got {ascii_font!r}"
        )


def test_iter_paragraph_runs_includes_hyperlink():
    """iter_paragraph_runs must yield runs inside w:hyperlink, not just direct runs."""
    from docx.oxml import OxmlElement
    from core.docx_utils import iter_paragraph_runs

    doc = Document()
    p = doc.add_paragraph("前文 ")

    # Add a hyperlink child with one run
    hyperlink = OxmlElement("w:hyperlink")
    r_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = "https://example.com"
    r_elem.append(t_elem)
    hyperlink.append(r_elem)
    p._p.append(hyperlink)

    # paragraph.runs only sees direct runs; iter_paragraph_runs should see all
    assert len(p.runs) == 1, "Baseline: paragraph.runs only sees the direct run, not the hyperlink run"
    all_runs = list(iter_paragraph_runs(p))
    assert len(all_runs) == 2, (
        f"iter_paragraph_runs should yield 2 runs (direct + hyperlink), got {len(all_runs)}"
    )
    texts = [r.text for r in all_runs]
    assert any(t == "https://example.com" for t in texts), "Hyperlink URL text not found in iter_paragraph_runs output"
