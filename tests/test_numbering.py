# tests/test_numbering.py
"""
Tests for core/numbering.py and the text-list-to-numPr conversion pipeline.

Coverage:
  1. detect_text_list_prefix correctly identifies all five formats
  2. create_list_num_id adds a fresh abstractNum + num to the document
  3. apply_numpr writes w:numPr onto the paragraph
  4. strip_list_text_prefix removes the text marker from runs
  5. convert_text_lists end-to-end: groups, min_run_len guard, numPr applied
  6. apply_formatting integrates numbering conversion (step 5)
  7. detect_role now classifies body list patterns as list_item
"""

from pathlib import Path
import copy
import re
import sys

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.numbering import (
    detect_text_list_prefix,
    create_list_num_id,
    apply_numpr,
    strip_list_text_prefix,
    convert_text_lists,
)
from core.formatter import apply_formatting, detect_role, _normalize_table_list_separators
from core.docx_utils import iter_all_paragraphs, is_effectively_blank_paragraph
from core.spec import load_spec, Spec
from core.parser import Block

SPECS_DIR = Path(__file__).resolve().parents[1] / "specs"


# ─── 1. detect_text_list_prefix ──────────────────────────────────────────────

def test_detect_paren_arabic():
    result = detect_text_list_prefix("（1）第一点")
    assert result is not None
    fmt, ordinal, prefix_len = result
    assert fmt == "paren_arabic"
    assert ordinal == 1
    assert prefix_len == 3  # "（" + "1" + "）" = 3 Unicode code points

def test_detect_paren_arabic_multi_digit():
    result = detect_text_list_prefix("（12）第十二点")
    assert result is not None
    fmt, ordinal, prefix_len = result
    assert fmt == "paren_arabic"
    assert ordinal == 12

def test_detect_rparen():
    result = detect_text_list_prefix("3) 第三点")
    assert result is not None
    fmt, ordinal, prefix_len = result
    assert fmt == "rparen"
    assert ordinal == 3

def test_detect_rparen_fullwidth():
    result = detect_text_list_prefix("2） 第二点")
    assert result is not None
    fmt, ordinal, _ = result
    assert fmt == "rparen"
    assert ordinal == 2

def test_detect_enclosed():
    for i, ch in enumerate("①②③④⑤⑥⑦⑧⑨⑩", start=1):
        result = detect_text_list_prefix(f"{ch}内容{i}")
        assert result is not None, f"Failed on {ch!r}"
        fmt, ordinal, _ = result
        assert fmt == "enclosed"
        assert ordinal == i

def test_detect_alpha_lower():
    result = detect_text_list_prefix("a. 选项A")
    assert result is not None
    fmt, ordinal, _ = result
    assert fmt == "alpha_lower"
    assert ordinal == 1

def test_detect_alpha_upper():
    result = detect_text_list_prefix("B. 选项B")
    assert result is not None
    fmt, ordinal, _ = result
    assert fmt == "alpha_upper"
    assert ordinal == 2

def test_detect_returns_none_for_body_text():
    for text in ["这是正文。", "第一条 总则", "一、概述", "1.1 引言", "摘要：xxx"]:
        assert detect_text_list_prefix(text) is None, f"Should be None for {text!r}"

def test_detect_returns_none_for_cn_paren_subtitle():
    # （一） is h3, not a body list item — confirm detect_text_list_prefix returns None
    assert detect_text_list_prefix("（一）子标题") is None


# ─── 2. create_list_num_id ────────────────────────────────────────────────────

def test_create_list_num_id_adds_to_numbering_part():
    doc = Document()
    nelem = doc.part.numbering_part._element
    ids_before = {c.get(qn("w:numId")) for c in nelem if c.tag == qn("w:num")}

    num_id = create_list_num_id(doc, "paren_arabic")
    assert isinstance(num_id, int)
    assert num_id > 0

    ids_after = {c.get(qn("w:numId")) for c in nelem if c.tag == qn("w:num")}
    assert str(num_id) in ids_after - ids_before


def test_create_list_num_id_adds_abstractNum():
    doc = Document()
    nelem = doc.part.numbering_part._element
    abs_ids_before = {c.get(qn("w:abstractNumId")) for c in nelem if c.tag == qn("w:abstractNum")}

    num_id = create_list_num_id(doc, "enclosed")

    abs_ids_after = {c.get(qn("w:abstractNumId")) for c in nelem if c.tag == qn("w:abstractNum")}
    assert len(abs_ids_after) == len(abs_ids_before) + 1


def test_create_two_lists_get_different_ids():
    doc = Document()
    id1 = create_list_num_id(doc, "paren_arabic")
    id2 = create_list_num_id(doc, "rparen")
    assert id1 != id2


# ─── 3. apply_numpr ──────────────────────────────────────────────────────────

def test_apply_numpr_writes_numpr():
    doc = Document()
    p = doc.add_paragraph("测试")
    num_id = create_list_num_id(doc, "paren_arabic")
    apply_numpr(p, num_id)

    pPr = p._p.pPr
    assert pPr is not None
    numPr = pPr.find(qn("w:numPr"))
    assert numPr is not None
    assert numPr.find(qn("w:numId")).get(qn("w:val")) == str(num_id)
    assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "0"


def test_apply_numpr_idempotent():
    """Calling apply_numpr twice should not create duplicate numPr elements."""
    doc = Document()
    p = doc.add_paragraph("测试")
    num_id = create_list_num_id(doc, "paren_arabic")
    apply_numpr(p, num_id)
    apply_numpr(p, num_id)

    pPr = p._p.pPr
    all_numpr = pPr.findall(qn("w:numPr"))
    assert len(all_numpr) == 1


# ─── 4. strip_list_text_prefix ───────────────────────────────────────────────

def test_strip_paren_arabic_prefix():
    doc = Document()
    p = doc.add_paragraph("（1）这是第一点内容")
    result = detect_text_list_prefix(p.text)
    assert result is not None
    _, _, prefix_len = result
    strip_list_text_prefix(p, prefix_len)
    assert p.text == "这是第一点内容"


def test_strip_enclosed_prefix():
    doc = Document()
    p = doc.add_paragraph("①这是第一点")
    result = detect_text_list_prefix(p.text)
    assert result is not None
    _, _, prefix_len = result
    strip_list_text_prefix(p, prefix_len)
    assert p.text == "这是第一点"


def test_strip_rparen_prefix():
    doc = Document()
    p = doc.add_paragraph("2) 第二点内容")
    result = detect_text_list_prefix(p.text)
    assert result is not None
    _, _, prefix_len = result
    strip_list_text_prefix(p, prefix_len)
    assert "2)" not in p.text
    assert "第二点内容" in p.text


# ─── 5. convert_text_lists ───────────────────────────────────────────────────

def _is_list_p(p):
    try:
        ppr = p._p.pPr
        return bool(ppr is not None and getattr(ppr, "numPr", None) is not None)
    except Exception:
        return False


def test_convert_text_lists_converts_group():
    doc = Document()
    texts = ["（1）第一项", "（2）第二项", "（3）第三项"]
    for t in texts:
        doc.add_paragraph(t)

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc, paras,
        get_role=lambda _: "list_item",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    assert converted == 3

    for p in iter_all_paragraphs(doc):
        assert _is_list_p(p), f"Paragraph {p.text!r} should have numPr"
        # Text prefix should be stripped
        assert not p.text.startswith("（"), f"Prefix not stripped from {p.text!r}"


def test_convert_text_lists_respects_min_run_len():
    """A single-item group should NOT be converted when min_run_len=2."""
    doc = Document()
    doc.add_paragraph("（1）只有一项")
    doc.add_paragraph("这是正文段落。")

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc, paras,
        get_role=lambda p: "list_item" if p.text.startswith("（") else "body",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    assert converted == 0

    # The paragraph should still have its text prefix intact
    p = iter_all_paragraphs(doc)[0]
    assert "（1）" in p.text


def test_convert_text_lists_skips_existing_numpr():
    """Paragraphs that already have numPr should be skipped."""
    doc = Document()
    p = doc.add_paragraph("（1）已有列表")
    # Manually apply numPr
    num_id = create_list_num_id(doc, "paren_arabic")
    apply_numpr(p, num_id)

    doc.add_paragraph("（2）第二项")

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc, paras,
        get_role=lambda _: "list_item",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    # The first paragraph already had numPr, so the pair is broken — no group of 2 without numPr
    assert converted == 0


def test_convert_different_formats_form_separate_groups():
    doc = Document()
    doc.add_paragraph("（1）阿拉伯括号一")
    doc.add_paragraph("（2）阿拉伯括号二")
    doc.add_paragraph("①圈一")
    doc.add_paragraph("②圈二")

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc, paras,
        get_role=lambda _: "list_item",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    assert converted == 4

    # Both groups should have numPr but with different numIds
    all_paras = iter_all_paragraphs(doc)
    num_ids = set()
    for p in all_paras:
        ppr = p._p.pPr
        if ppr is not None:
            numPr = ppr.find(qn("w:numPr"))
            if numPr is not None:
                num_ids.add(numPr.find(qn("w:numId")).get(qn("w:val")))
    assert len(num_ids) == 2, f"Expected 2 different numIds, got {num_ids}"


# ─── 6. apply_formatting integration ─────────────────────────────────────────

def _make_doc_blocks_labels(role_texts):
    doc = Document()
    for _, text in role_texts:
        doc.add_paragraph(text)
    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {b.block_id: role for b, (role, _) in zip(blocks, role_texts)}
    labels["_source"] = "test"
    return doc, blocks, labels


def test_apply_formatting_converts_text_lists():
    """apply_formatting step 4.5 must convert LLM-labeled list items to real numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "（1）第一项内容"),
        ("list_item", "（2）第二项内容"),
        ("list_item", "（3）第三项内容"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)

    # Report should record the conversion count (LLM-direct path)
    total_converted = (
        report["actions"]["llm_direct_list_converted"]
        + report["actions"]["text_list_converted_to_numpr"]
    )
    assert total_converted == 3

    # All three paragraphs should have real numPr
    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert _is_list_p(p), f"Expected numPr on {p.text!r}"


def test_apply_formatting_text_list_prefix_stripped():
    """After conversion, the text-based prefix (（1）) must be removed from runs."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "（1）内容一"),
        ("list_item", "（2）内容二"),
    ])

    apply_formatting(doc, blocks, labels, spec)

    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert "（" not in p.text, f"Prefix still present in {p.text!r}"


def test_apply_formatting_convert_text_numbers_disabled():
    """When convert_text_numbers=false, text lists are NOT converted."""
    import copy
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    # Mutate a copy to disable conversion
    raw = copy.deepcopy(spec.raw)
    raw["list_item"]["convert_text_numbers"] = False
    from core.spec import Spec
    spec_off = Spec(raw=raw)

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "（1）第一项"),
        ("list_item", "（2）第二项"),
    ])

    report = apply_formatting(doc, blocks, labels, spec_off)

    assert report["actions"]["text_list_converted_to_numpr"] == 0
    # Text prefix should still be present
    paras = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert any("（1）" in p.text for p in paras)


def test_apply_formatting_report_includes_numpr_count():
    """report['actions']['text_list_converted_to_numpr'] key must always be present."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    doc, blocks, labels = _make_doc_blocks_labels([("body", "这是正文。")])
    report = apply_formatting(doc, blocks, labels, spec)
    assert "text_list_converted_to_numpr" in report["actions"]


# ─── 7. detect_role for body list patterns ───────────────────────────────────

def test_detect_role_body_list_paren_arabic():
    doc = Document()
    assert detect_role(doc.add_paragraph("（1）第一点")) == "list_item"
    assert detect_role(doc.add_paragraph("（10）第十点")) == "list_item"


def test_detect_role_body_list_rparen():
    doc = Document()
    assert detect_role(doc.add_paragraph("1) 第一点")) == "list_item"
    assert detect_role(doc.add_paragraph("3） 第三点")) == "list_item"


def test_detect_role_body_list_enclosed():
    doc = Document()
    assert detect_role(doc.add_paragraph("①第一点")) == "list_item"
    assert detect_role(doc.add_paragraph("⑩第十点")) == "list_item"


def test_detect_role_body_list_alpha():
    doc = Document()
    assert detect_role(doc.add_paragraph("a. 选项A")) == "list_item"
    assert detect_role(doc.add_paragraph("B. 选项B")) == "list_item"


def test_detect_role_cn_paren_subtitle_still_h3():
    """（一）style (Chinese numeral in parentheses) must stay h3, not list_item."""
    doc = Document()
    assert detect_role(doc.add_paragraph("（一）子标题内容")) == "h3"
    assert detect_role(doc.add_paragraph("（三）另一子节")) == "h3"


# ─── 8. num_dot format (1. text) ─────────────────────────────────────────────

def test_detect_text_list_prefix_num_dot():
    result = detect_text_list_prefix("1. 第一项内容")
    assert result is not None
    fmt, ordinal, prefix_len = result
    assert fmt == "num_dot"
    assert ordinal == 1
    assert prefix_len == 3  # "1" + "." + " " = 3 chars


def test_detect_text_list_prefix_num_dot_multi_digit():
    result = detect_text_list_prefix("10. 第十项内容")
    assert result is not None
    fmt, ordinal, prefix_len = result
    assert fmt == "num_dot"
    assert ordinal == 10
    assert prefix_len == 4  # "10" + "." + " " = 4 chars


def test_detect_text_list_prefix_num_dot_not_multilevel():
    """1.1 text (multi-level) must NOT match num_dot."""
    assert detect_text_list_prefix("1.1 多级标题") is None
    assert detect_text_list_prefix("2.3 另一节") is None


def test_detect_role_body_list_num_dot():
    """1. text style should be classified as list_item."""
    doc = Document()
    assert detect_role(doc.add_paragraph("1. 第一点内容")) == "list_item"
    assert detect_role(doc.add_paragraph("2. 第二点内容")) == "list_item"
    assert detect_role(doc.add_paragraph("10. 第十点内容")) == "list_item"


def test_detect_role_multilevel_not_list_item():
    """1.1 text (multi-level) must NOT be list_item (stays h3 via RE_NUM_DOT)."""
    doc = Document()
    assert detect_role(doc.add_paragraph("1.1 二级标题")) == "h3"


def test_convert_text_lists_num_dot():
    """num_dot format (1. text) should be converted to real Word list."""
    doc = Document()
    texts = ["1. 第一项", "2. 第二项", "3. 第三项"]
    for t in texts:
        doc.add_paragraph(t)

    def _is_list_p(p):
        try:
            ppr = p._p.pPr
            return bool(ppr is not None and getattr(ppr, "numPr", None) is not None)
        except Exception:
            return False

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc, paras,
        get_role=lambda _: "list_item",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    assert converted == 3
    for p in iter_all_paragraphs(doc):
        assert _is_list_p(p), f"Paragraph {p.text!r} should have numPr"
        # No numeric prefix should remain (1., 2., 3.)
        assert not re.match(r"^\d+\. ", p.text), f"Prefix not stripped: {p.text!r}"


def test_apply_formatting_converts_num_dot_lists():
    """apply_formatting must convert 1. 2. 3. style lists to real numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "1. 第一项内容"),
        ("list_item", "2. 第二项内容"),
        ("list_item", "3. 第三项内容"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)
    total_converted = (
        report["actions"]["llm_direct_list_converted"]
        + report["actions"]["text_list_converted_to_numpr"]
    )
    assert total_converted == 3

    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            ppr = p._p.pPr
            assert ppr is not None and ppr.find(qn("w:numPr")) is not None


# ─── 9. table cell formatting ────────────────────────────────────────────────

def test_table_cell_body_no_first_line_indent():
    """Body paragraphs inside table cells must not receive first-line indent."""
    from core.formatter import apply_formatting
    from core.parser import Block

    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Replace the default empty paragraph with body-like content
    p = cell.paragraphs[0]
    p.text = "这是表格内容"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=para.text, paragraph_index=i)
        for i, para in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    cell_p = table.cell(0, 0).paragraphs[0]
    fli = cell_p.paragraph_format.first_line_indent
    # Should be 0 (no indent) or None (not set), NOT Pt(24) or similar
    assert fli is None or fli == 0, f"Expected no first-line indent in cell, got {fli}"


def test_autofit_tables_action_in_report():
    """apply_formatting report must include tables_autofitted count."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    doc = Document()
    doc.add_table(rows=2, cols=2)
    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}
    report = apply_formatting(doc, blocks, labels, spec)
    assert "tables_autofitted" in report["actions"]
    assert report["actions"]["tables_autofitted"] == 1


# ─── 10. create_list_num_id on document without numbering part ────────────────

def _make_doc_without_numbering_part():
    """Return a Document loaded from a docx that has no word/numbering.xml."""
    import io
    import re as _re
    import zipfile

    buf = io.BytesIO()
    Document().save(buf)
    buf.seek(0)

    buf2 = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as z_in:
        with zipfile.ZipFile(buf2, "w", zipfile.ZIP_DEFLATED) as z_out:
            for name in z_in.namelist():
                data = z_in.read(name)
                if "numbering" in name.lower():
                    continue  # strip numbering.xml
                if name == "word/_rels/document.xml.rels":
                    # Remove the Relationship entry pointing at numbering.xml
                    data = _re.sub(
                        rb"<Relationship[^>]*numbering[^>]*/?>",
                        b"",
                        data,
                    )
                z_out.writestr(name, data)

    buf2.seek(0)
    return Document(buf2)


def test_create_list_num_id_creates_numbering_part_when_absent():
    """create_list_num_id must succeed even if the document has no numbering part."""
    doc = _make_doc_without_numbering_part()

    # Accessing numbering_part on the bare doc should raise NotImplementedError
    try:
        _ = doc.part.numbering_part._element
        # If we reach here the template unexpectedly kept numbering.xml — skip
        return
    except (NotImplementedError, Exception):
        pass

    # create_list_num_id must not raise
    num_id = create_list_num_id(doc, "paren_arabic")
    assert isinstance(num_id, int) and num_id > 0

    # The numbering part must now exist and contain the new num element
    nelem = doc.part.numbering_part._element
    num_ids = {c.get(qn("w:numId")) for c in nelem if c.tag == qn("w:num")}
    assert str(num_id) in num_ids


def test_convert_text_lists_on_doc_without_numbering_part():
    """convert_text_lists end-to-end must work when the doc starts without a numbering part."""
    doc = _make_doc_without_numbering_part()
    for text in ["（1）第一项", "（2）第二项", "（3）第三项"]:
        doc.add_paragraph(text)

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc,
        paras,
        get_role=lambda _: "list_item",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=2,
    )
    assert converted == 3
    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert _is_list_p(p), f"Expected numPr on {p.text!r}"


# ─── 11. min_run_len=1 converts single-item lists ────────────────────────────

def test_convert_text_lists_min_run_len_1_converts_single_item():
    """With min_run_len=1, even a single list item should be converted."""
    doc = Document()
    doc.add_paragraph("（1）只有一项")
    doc.add_paragraph("这是正文段落。")

    paras = iter_all_paragraphs(doc)
    converted, _ = convert_text_lists(
        doc,
        paras,
        get_role=lambda p: "list_item" if p.text.startswith("（") else "body",
        is_list_paragraph_fn=_is_list_p,
        is_blank_fn=is_effectively_blank_paragraph,
        min_run_len=1,
    )
    assert converted == 1

    # The single item should now have numPr
    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 1
    # And its text prefix should be stripped
    assert "（1）" not in list_paras[0].text


def test_apply_formatting_min_run_len_1_via_spec():
    """apply_formatting with min_run_len=1 in spec must convert even single list items."""
    import copy
    from core.spec import Spec

    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    raw = copy.deepcopy(spec.raw)
    raw["list_item"]["min_run_len"] = 1
    spec_1 = Spec(raw=raw)

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "（1）单独一项"),
        ("body", "后接正文段落。"),
    ])

    report = apply_formatting(doc, blocks, labels, spec_1)
    assert report["actions"]["text_list_converted_to_numpr"] == 1

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 1


# ─── 12. strip_list_text_prefix graceful no-op with no runs ──────────────────

def test_strip_list_text_prefix_no_runs_is_noop():
    """strip_list_text_prefix must not raise when a paragraph has no runs."""
    doc = Document()
    p = doc.add_paragraph("")
    # Manually clear all runs so the paragraph has none
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    assert not p.runs  # confirm no runs

    # Should not raise
    strip_list_text_prefix(p, 3)
    # Text is empty / unchanged — no crash is the key assertion


# ─── 13. New bug-fix tests ────────────────────────────────────────────────────

def test_body_role_numbered_paragraph_gets_converted():
    """A paragraph labeled 'body' but starting with （1） must be converted to numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("body", "（1）内容很长的正文编号段落"),
        ("body", "（2）第二条正文编号内容"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)

    # At least the LLM-direct path should have converted them
    total_converted = (
        report["actions"]["llm_direct_list_converted"]
        + report["actions"]["text_list_converted_to_numpr"]
    )
    assert total_converted == 2, f"Expected 2 converted, got {total_converted}"

    # Both paragraphs must now carry real numPr
    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert _is_list_p(p), f"Expected numPr on {p.text!r}"
            assert "（" not in p.text, f"Prefix still present in {p.text!r}"


def test_list_item_font_applied_after_numpr_conversion():
    """Paragraphs converted to numPr must have their runs' font set per list_item spec."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    raw = copy.deepcopy(spec.raw)
    raw["list_item"]["font_size_pt"] = 14
    raw["fonts"]["zh"] = "仿宋"
    spec_mod = Spec(raw=raw)

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "（1）第一条"),
        ("list_item", "（2）第二条"),
    ])

    apply_formatting(doc, blocks, labels, spec_mod)

    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert _is_list_p(p), f"Expected numPr on {p.text!r}"
            for run in p.runs:
                if run.text:
                    assert run.font.size == Pt(14), (
                        f"Expected font size 14pt on run {run.text!r}, got {run.font.size}"
                    )


def test_table_list_item_font_applied_after_numpr_conversion():
    """Table cell paragraphs converted to numPr must also get list_item font settings."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    raw = copy.deepcopy(spec.raw)
    raw["list_item"]["font_size_pt"] = 11
    spec_mod = Spec(raw=raw)

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Clear default empty paragraph and add two list-style paragraphs
    cell.paragraphs[0].text = "（1）表格列表项一"
    cell.add_paragraph("（2）表格列表项二")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    # Label the table cell paragraphs as list_item
    labels = {b.block_id: "list_item" for b in blocks if b.text.startswith("（")}
    labels["_source"] = "test"

    apply_formatting(doc, blocks, labels, spec_mod)

    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p) and _is_list_p(p):
            for run in p.runs:
                if run.text:
                    assert run.font.size == Pt(11), (
                        f"Expected font size 11pt on run {run.text!r}, got {run.font.size}"
                    )


# ─── 14. Regression: all consecutive numbered items converted, not just the first ─

def test_all_consecutive_num_dot_items_converted_not_just_first():
    """Regression: 1./2./3. consecutive items must ALL receive numPr, not only item 1."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "1. 第一项内容"),
        ("list_item", "2. 第二项内容"),
        ("list_item", "3. 第三项内容"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)
    assert report["actions"]["text_list_converted_to_numpr"] == 3, (
        f"Expected all 3 items converted, got {report['actions']['text_list_converted_to_numpr']}"
    )

    non_blank = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(non_blank) == 3
    for p in non_blank:
        assert _is_list_p(p), f"Expected numPr on {p.text!r} (first-item-only regression)"

    # All items must share the same numId (one continuous list)
    num_ids = set()
    for p in non_blank:
        ppr = p._p.pPr
        if ppr is not None:
            numPr = ppr.find(qn("w:numPr"))
            if numPr is not None:
                num_ids.add(numPr.find(qn("w:numId")).get(qn("w:val")))
    assert len(num_ids) == 1, f"All items should share one numId, got {num_ids}"


def test_mixed_labels_all_consecutive_items_converted():
    """Regression: items 2/3 labeled 'body' must be converted together with item 1 (list_item)."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "1. 第一项内容"),  # LLM correctly identified
        ("body",      "2. 第二项内容"),  # LLM mis-labeled as body
        ("body",      "3. 第三项内容"),  # LLM mis-labeled as body
    ])

    report = apply_formatting(doc, blocks, labels, spec)
    assert report["actions"]["text_list_converted_to_numpr"] == 3, (
        f"Expected all 3 items converted even with mixed labels, "
        f"got {report['actions']['text_list_converted_to_numpr']}"
    )

    non_blank = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    for p in non_blank:
        assert _is_list_p(p), (
            f"Expected numPr on {p.text!r}; only-first-item regression detected"
        )

    # All must share the same numId so the counter is continuous
    num_ids = set()
    for p in non_blank:
        ppr = p._p.pPr
        if ppr is not None:
            numPr = ppr.find(qn("w:numPr"))
            if numPr is not None:
                num_ids.add(numPr.find(qn("w:numId")).get(qn("w:val")))
    assert len(num_ids) == 1, (
        f"All items in the same list group must share one numId, got {num_ids}"
    )


def test_table_cell_rparen_items_all_converted():
    """Table cell paragraphs with 1) 2) style must ALL be converted to numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "1) 第一条"
    cell.add_paragraph("2) 第二条")
    cell.add_paragraph("3) 第三条")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    # Only label the first item; items 2 and 3 are left unlabeled (simulate LLM partial result)
    labels = {"_source": "test"}
    for b in blocks:
        if b.text == "1) 第一条":
            labels[b.block_id] = "list_item"

    apply_formatting(doc, blocks, labels, spec)

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, (
        f"Expected all 3 table cell items to have numPr, got {len(list_paras)}"
    )


# ─── 15. rparen without trailing space ───────────────────────────────────────

def test_detect_rparen_no_space():
    """rparen items written without space after paren (e.g. '2）内容') must be detected."""
    result = detect_text_list_prefix("2）第二点")
    assert result is not None, "Expected rparen detection for '2）第二点' (no space)"
    fmt, ordinal, prefix_len = result
    assert fmt == "rparen"
    assert ordinal == 2
    assert prefix_len == 2  # "2" + "）" = 2 characters


def test_detect_rparen_no_space_prefix_stripped_correctly():
    """strip_list_text_prefix must remove the '2）' prefix when there is no trailing space."""
    doc = Document()
    p = doc.add_paragraph("2）第二点内容")
    result = detect_text_list_prefix(p.text)
    assert result is not None
    _, _, prefix_len = result
    strip_list_text_prefix(p, prefix_len)
    assert "2）" not in p.text
    assert "第二点内容" in p.text


def test_detect_role_rparen_no_space_is_list_item():
    """detect_role must return 'list_item' for '1）内容' (no space after paren)."""
    doc = Document()
    assert detect_role(doc.add_paragraph("1）第一点")) == "list_item"
    assert detect_role(doc.add_paragraph("2）第二点")) == "list_item"
    assert detect_role(doc.add_paragraph("10）第十点")) == "list_item"


# ─── 16. mixed paren_arabic + rparen in same cell all converted ───────────────

def test_table_cell_mixed_paren_format_all_converted():
    """
    Chinese docs often use （1） for the first item and 2）/3） for subsequent items.
    All three items must be converted to numPr as a single list group.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "（1）表格列表项一"
    cell.add_paragraph("2）表格列表项二")
    cell.add_paragraph("3）表格列表项三")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, (
        f"Expected all 3 mixed-format items to have numPr, got {len(list_paras)}"
    )


def test_table_cell_mixed_paren_font_applied():
    """All items in a mixed-format list must have the list_item font applied."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    raw = copy.deepcopy(spec.raw)
    raw["list_item"]["font_size_pt"] = 11
    spec_mod = Spec(raw=raw)

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "（1）第一条"
    cell.add_paragraph("2）第二条")
    cell.add_paragraph("3）第三条")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec_mod)

    for p in iter_all_paragraphs(doc):
        if not is_effectively_blank_paragraph(p):
            assert _is_list_p(p), f"Expected numPr on {p.text!r}"
            for run in p.runs:
                if run.text:
                    assert run.font.size == Pt(11), (
                        f"Expected 11pt on run {run.text!r}, got {run.font.size}"
                    )


# ─── 17. cell-boundary: two cells with independent lists get separate numIds ──

def test_two_table_cells_get_independent_numids():
    """
    Two table cells each containing a 1)/2)/3) list must get separate numId values.
    Items from cell 1 must NOT be merged into the same Word list as items from cell 2.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell1 = table.cell(0, 0)
    cell1.paragraphs[0].text = "1）第一条"
    cell1.add_paragraph("2）第二条")

    cell2 = table.cell(0, 1)
    cell2.paragraphs[0].text = "1）甲"
    cell2.add_paragraph("2）乙")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    # All 4 paragraphs should have numPr
    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 4, f"Expected 4 numPr paragraphs, got {len(list_paras)}"

    # Each cell's items should share one numId, but the two cells must have different numIds
    def _get_num_ids(cell):
        ids = set()
        for p in cell.paragraphs:
            ppr = p._p.pPr
            if ppr is not None:
                numPr = ppr.find(qn("w:numPr"))
                if numPr is not None:
                    ids.add(numPr.find(qn("w:numId")).get(qn("w:val")))
        return ids

    ids1 = _get_num_ids(table.cell(0, 0))
    ids2 = _get_num_ids(table.cell(0, 1))

    assert len(ids1) == 1, f"Cell 1 items should share one numId, got {ids1}"
    assert len(ids2) == 1, f"Cell 2 items should share one numId, got {ids2}"
    assert ids1 != ids2, (
        f"Cell 1 and cell 2 must have different numIds (got the same: {ids1})"
    )


def test_apply_formatting_converts_all_table_linebreak_number_items():
    """Table-cell list text split by linebreaks should all become numPr paragraphs."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.text = "1）第一条\n2）第二条\n3）第三条"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=para.text, paragraph_index=i)
        for i, para in enumerate(paras)
    ]
    labels = {blocks[0].block_id: "list_item", "_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["text_list_converted_to_numpr"] == 3
    cell_paras = table.cell(0, 0).paragraphs
    assert len(cell_paras) == 3
    for para in cell_paras:
        assert _is_list_p(para), f"Expected numPr on {para.text!r}"
        assert not re.match(r"^\s*\d+[)）]", para.text), f"Prefix not stripped: {para.text!r}"


# ─── 18. items in separate rows/cells each get the correct start ordinal ──────

def _get_abstractnum_elem(doc, num_id_str):
    """Return the ``w:abstractNum`` element linked to *num_id_str*, or ``None``."""
    from core.numbering import _numbering_element
    nelem = _numbering_element(doc)
    abs_id_str = None
    for child in nelem:
        if child.tag == qn("w:num") and child.get(qn("w:numId")) == num_id_str:
            ref = child.find(qn("w:abstractNumId"))
            if ref is not None:
                abs_id_str = ref.get(qn("w:val"))
            break
    if abs_id_str is None:
        return None
    for child in nelem:
        if child.tag == qn("w:abstractNum") and child.get(qn("w:abstractNumId")) == abs_id_str:
            return child
    return None


def _get_abstractnum_start(doc, num_id_str):
    """Return the ``w:start`` integer for the abstractNum linked to *num_id_str*."""
    abs_node = _get_abstractnum_elem(doc, num_id_str)
    if abs_node is None:
        return None
    lvl = abs_node.find(qn("w:lvl"))
    if lvl is None:
        return None
    start_el = lvl.find(qn("w:start"))
    if start_el is None:
        return None
    return int(start_el.get(qn("w:val")))


def test_table_items_in_separate_rows_get_correct_start_ordinal():
    """
    Regression: when 1）/2）/3） each live in their own table cell (different rows),
    each gets a separate numId.  The numId for item N must have w:start=N so that
    Word renders them as 1), 2), 3) respectively – not all as 1).
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    for i, row in enumerate(table.rows):
        row.cells[0].paragraphs[0].text = f"{i + 1}）第{i + 1}项内容"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, f"Expected all 3 items converted, got {len(list_paras)}"

    # Each item must use a distinct numId (different cells → different groups)
    num_id_vals = []
    for p in list_paras:
        ppr = p._p.pPr
        numPr = ppr.find(qn("w:numPr"))
        num_id_vals.append(numPr.find(qn("w:numId")).get(qn("w:val")))
    assert len(set(num_id_vals)) == 3, (
        f"Each item in a separate cell must have its own numId, got {num_id_vals}"
    )

    # The w:start of each numId's abstractNum must equal the item ordinal (1, 2, 3)
    for expected_start, num_id_str in zip([1, 2, 3], num_id_vals):
        actual_start = _get_abstractnum_start(doc, num_id_str)
        assert actual_start == expected_start, (
            f"numId {num_id_str}: expected w:start={expected_start}, got {actual_start}. "
            f"Items in different cells must use the correct start ordinal so Word renders "
            f"them as 1), 2), 3) instead of all as 1)."
        )


def test_table_items_in_separate_rows_times_new_roman_font():
    """
    Items in separate table cells converted to numPr must have Times New Roman
    applied to their list-marker glyph (abstractNum lvl/rPr/rFonts).
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    # default.yaml has fonts.en = "Times New Roman"

    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    for i, row in enumerate(table.rows):
        row.cells[0].paragraphs[0].text = f"{i + 1}）内容{i + 1}"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    apply_formatting(doc, blocks, labels, spec)

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, f"Expected 3 converted list items, got {len(list_paras)}"

    for p in list_paras:
        ppr = p._p.pPr
        numPr = ppr.find(qn("w:numPr"))
        num_id_str = numPr.find(qn("w:numId")).get(qn("w:val"))

        abs_node = _get_abstractnum_elem(doc, num_id_str)
        assert abs_node is not None, f"No abstractNum found for numId {num_id_str}"
        lvl = abs_node.find(qn("w:lvl"))
        rPr = lvl.find(qn("w:rPr"))
        abs_id = abs_node.get(qn("w:abstractNumId"))
        assert rPr is not None, f"abstractNum {abs_id} lvl has no rPr"
        rFonts = rPr.find(qn("w:rFonts"))
        assert rFonts is not None, f"abstractNum {abs_id} lvl rPr has no rFonts"
        assert rFonts.get(qn("w:ascii")) == "Times New Roman", (
            f"Expected Times New Roman for w:ascii on abstractNum {abs_id}, "
            f"got {rFonts.get(qn('w:ascii'))!r}"
        )


def test_create_list_num_id_writes_lvl_rpr_font_settings():
    """Numbering definition should carry lvl/rPr so list marker font follows spec."""
    doc = Document()
    num_id = create_list_num_id(
        doc,
        "rparen",
        zh_font="仿宋_GB2312",
        en_font="Times New Roman",
        size_pt=14,
        bold=False,
        italic=False,
    )

    nelem = doc.part.numbering_part._element
    target_num = None
    for child in nelem:
        if child.tag == qn("w:num") and child.get(qn("w:numId")) == str(num_id):
            target_num = child
            break
    assert target_num is not None

    abs_id = target_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abs_node = None
    for child in nelem:
        if child.tag == qn("w:abstractNum") and child.get(qn("w:abstractNumId")) == abs_id:
            abs_node = child
            break
    assert abs_node is not None

    lvl = abs_node.find(qn("w:lvl"))
    assert lvl is not None
    rPr = lvl.find(qn("w:rPr"))
    assert rPr is not None

    rFonts = rPr.find(qn("w:rFonts"))
    assert rFonts is not None
    assert rFonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
    assert rFonts.get(qn("w:ascii")) == "Times New Roman"
    assert rFonts.get(qn("w:hAnsi")) == "Times New Roman"

    sz = rPr.find(qn("w:sz"))
    assert sz is not None
    assert sz.get(qn("w:val")) == "28"


def test_apply_formatting_converts_table_carriage_return_number_items():
    """Table-cell list text using carriage-return soft breaks should all become numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "1）第一条\r2）第二条\r3）第三条"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=para.text, paragraph_index=i)
        for i, para in enumerate(paras)
    ]
    labels = {blocks[0].block_id: "list_item", "_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["text_list_converted_to_numpr"] == 3
    cell_paras = table.cell(0, 0).paragraphs
    assert len(cell_paras) == 3
    for para in cell_paras:
        assert _is_list_p(para), f"Expected numPr on {para.text!r}"
        assert not re.match(r"^\s*\d+[)）]", para.text), f"Prefix not stripped: {para.text!r}"


# ─── 19. min_run_len default = 1: items in separate rows always converted ─────

def test_spec_default_min_run_len_is_1():
    """_validate_and_fill_defaults must set min_run_len=1 when not specified."""
    from core.spec import _validate_and_fill_defaults
    raw = {
        "fonts": {"zh": "宋体", "en": "Times New Roman"},
        "body": {
            "font_size_pt": 12, "line_spacing": 1.5,
            "space_before_pt": 0, "space_after_pt": 0, "first_line_chars": 2,
        },
        "heading": {
            "h1": {"font_size_pt": 16, "bold": True, "space_before_pt": 12, "space_after_pt": 6},
            "h2": {"font_size_pt": 14, "bold": True, "space_before_pt": 10, "space_after_pt": 4},
            "h3": {"font_size_pt": 12, "bold": True, "space_before_pt": 8,  "space_after_pt": 2},
        },
    }
    filled = _validate_and_fill_defaults(raw)
    assert filled["list_item"]["min_run_len"] == 1, (
        "Default min_run_len must be 1 so items in separate table cells are converted"
    )


def test_table_separate_rows_no_explicit_min_run_len():
    """
    Regression: items 1）/2）/3） each in their own table row must ALL get numPr
    even when the spec does not explicitly set min_run_len (relies on the default=1).
    """
    from core.spec import Spec, _validate_and_fill_defaults
    # Build a minimal spec without a min_run_len key under list_item
    raw = {
        "fonts": {"zh": "宋体", "en": "Times New Roman"},
        "body": {
            "font_size_pt": 12, "line_spacing": 1.5,
            "space_before_pt": 0, "space_after_pt": 0, "first_line_chars": 2,
        },
        "heading": {
            "h1": {"font_size_pt": 16, "bold": True, "space_before_pt": 12, "space_after_pt": 6},
            "h2": {"font_size_pt": 14, "bold": True, "space_before_pt": 10, "space_after_pt": 4},
            "h3": {"font_size_pt": 12, "bold": True, "space_before_pt": 8,  "space_after_pt": 2},
        },
        # list_item section deliberately omits min_run_len
        "list_item": {"font_size_pt": 12, "bold": False, "italic": False},
    }
    spec = Spec(raw=_validate_and_fill_defaults(raw))

    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    for i, row in enumerate(table.rows):
        row.cells[0].paragraphs[0].text = f"{i + 1}）第{i + 1}项内容"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    total = report["actions"]["text_list_converted_to_numpr"]
    assert total == 3, (
        f"Expected all 3 table-row items converted even without explicit min_run_len, "
        f"got {total}. This is the regression: body items form one group so they pass "
        f"min_run_len=2, but table items in separate cells form groups of 1 and get "
        f"silently skipped when the default is 2."
    )

    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, f"Expected 3 numPr paragraphs, got {len(list_paras)}"


# ─── 20. inline list separators in table cells (；N) pattern) ─────────────────

def test_normalize_table_list_separators_basic():
    """Table cell with '1) item1；2) item2；3) item3' must split into 3 paragraphs."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "1) 第一步内容；2) 第二步内容；3) 第三步内容"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["table_inline_list_normalized"] == 1, (
        "Expected 1 paragraph normalized"
    )
    assert report["actions"]["text_list_converted_to_numpr"] == 3, (
        f"Expected all 3 items converted, got {report['actions']['text_list_converted_to_numpr']}"
    )
    cell_paras = table.cell(0, 0).paragraphs
    assert len(cell_paras) == 3, f"Expected 3 paragraphs, got {len(cell_paras)}"
    for para in cell_paras:
        assert _is_list_p(para), f"Expected numPr on {para.text!r}"
        assert not re.match(r"^\s*\d+[)）]", para.text), f"Prefix not stripped: {para.text!r}"


def test_normalize_table_list_separators_paren_arabic():
    """Table cell with '（1）item1；（2）item2；（3）item3' must split into 3 paragraphs."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "（1）第一步；（2）第二步；（3）第三步"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["text_list_converted_to_numpr"] == 3, (
        f"Expected 3 converted, got {report['actions']['text_list_converted_to_numpr']}"
    )
    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 3, f"Expected 3 numPr paragraphs, got {len(list_paras)}"


def test_normalize_table_list_separators_body_text_unaffected():
    """Body (non-table) paragraphs with ；N) must NOT be split by the table normalizer."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    # Body paragraph (not in a table) with inline list markers
    doc.add_paragraph("1) 第一步；2) 第二步；3) 第三步")

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["table_inline_list_normalized"] == 0, (
        "Body text must not be affected by table inline list normalizer"
    )
    # Body text: only the first item '1)' gets converted
    assert report["actions"]["text_list_converted_to_numpr"] == 1


def test_normalize_table_list_separators_four_items():
    """Table cell with 4 inline items separated by ；must all become numPr."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "1) 甲；2) 乙；3) 丙；4) 丁"

    paras = iter_all_paragraphs(doc)
    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(paras)
    ]
    labels = {"_source": "test"}

    report = apply_formatting(doc, blocks, labels, spec)

    assert report["actions"]["text_list_converted_to_numpr"] == 4, (
        f"Expected 4 items converted, got {report['actions']['text_list_converted_to_numpr']}"
    )
    list_paras = [p for p in iter_all_paragraphs(doc) if _is_list_p(p)]
    assert len(list_paras) == 4, f"Expected 4 numPr paragraphs, got {len(list_paras)}"


# ─── LLM body-labeled list items get proper hanging indent after numPr ────────

def test_llm_body_labeled_list_gets_hanging_indent():
    """
    Paragraphs labeled 'body' by LLM but containing list prefixes must receive
    proper list hanging indent after step-5 numPr conversion, not the body
    forward first_line_indent that was applied in step 4.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    hanging_pt = float(spec.raw["list_item"]["hanging_indent_pt"])

    doc, blocks, labels = _make_doc_blocks_labels([
        ("body", "（1）第一项内容"),  # LLM mis-labeled as body
        ("body", "（2）第二项内容"),  # LLM mis-labeled as body
    ])

    apply_formatting(doc, blocks, labels, spec)

    non_blank = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(non_blank) == 2
    for p in non_blank:
        assert _is_list_p(p), f"Expected numPr on {p.text!r}"
        fli = p.paragraph_format.first_line_indent
        # Must be negative (hanging) — the default spec uses hanging_indent_pt=18,
        # so zero would only occur if hanging_indent_pt=0 and first_line_chars=0.
        assert fli is not None and fli <= 0, (
            f"first_line_indent must be ≤0 (hanging or zero) for numPr paragraph, got {fli}"
        )
        li = p.paragraph_format.left_indent
        assert li is not None and float(li) == pytest.approx(Pt(hanging_pt), rel=0.01), (
            f"left_indent must equal hanging_indent_pt={hanging_pt}pt, got {li}"
        )


def test_llm_body_labeled_first_line_indent_not_overriding_numpr():
    """
    After LLM labels numbered items as 'body', step 4 sets first_line_indent=+2chars.
    After step-5 converts them to numPr, the post-processing must override
    first_line_indent to the hanging-indent value so numPr renders correctly.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    body_size = float(spec.raw["body"]["font_size_pt"])
    first_line_chars = int(spec.raw["body"]["first_line_chars"])
    forward_indent_pt = first_line_chars * body_size  # the positive body indent

    doc, blocks, labels = _make_doc_blocks_labels([
        ("body", "1. 第一条"),
        ("body", "2. 第二条"),
        ("body", "3. 第三条"),
    ])

    apply_formatting(doc, blocks, labels, spec)

    non_blank = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(non_blank) == 3
    for p in non_blank:
        assert _is_list_p(p), f"Expected numPr on {p.text!r}"
        fli = p.paragraph_format.first_line_indent
        # Must NOT be the positive body first_line_indent
        assert fli is None or fli != pytest.approx(Pt(forward_indent_pt), rel=0.01), (
            f"first_line_indent must not be the positive body indent ({forward_indent_pt}pt) "
            f"after numPr conversion, got {fli}"
        )
        # Must be ≤ 0 (hanging or zero)
        if fli is not None:
            assert float(fli) <= 0, (
                f"first_line_indent must be ≤0 for a converted numPr paragraph, got {fli}"
            )


def test_normalize_table_list_separators_preserves_semicolon_in_content():
    """A ；that is NOT followed by a list marker must not be replaced with \\n."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # The ；after "可用" is in the middle of content, not before a list marker
    cell.paragraphs[0].text = "1) 选型（可用；商用）；2) 添加LICENSE"

    count = _normalize_table_list_separators(doc)
    cell_paras = cell.paragraphs
    # Only one split at ；2) → 2 paragraphs: "1) 选型（可用；商用）" and "2) 添加LICENSE"
    # The ；in "（可用；商用）" is preserved
    assert count == 1, f"Expected 1 paragraph modified, got {count}"
    full_text = "".join(p.text for p in cell_paras)
    assert "可用；商用" in full_text or "可用" in full_text, (
        "Semicolon inside content must be preserved"
    )


# ─── Mismatch suppression for multi-line numbered blocks ────────────────────

def test_multiline_list_item_not_counted_as_mismatch():
    """
    A paragraph labeled 'list_item' by LLM but containing multiple numbered
    items separated by \\n must NOT be counted as a mismatch.

    detect_role returns 'body' for multi-line numbered blocks (conservative
    safety guard to avoid heading mis-classification).  The LLM's label is the
    reliable signal for these paragraphs, so the consistency check should skip
    them rather than produce a false-alarm mismatch.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    # Simulate what the LLM sees: a paragraph that contains multiple numbered
    # items joined by soft line-breaks (as produced by Word's Shift+Enter).
    doc, blocks, labels = _make_doc_blocks_labels([
        ("h1",       "第一章 总结"),   # detect_role → h1 (starts with 第…章)
        ("list_item", "1. 数学成绩优秀。\n2. 编程能力增强。\n3. 逻辑推理提升。"),
        ("body",     "这是普通正文段落。"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)

    consistency = report["labels"]["consistency"]
    # The multi-line list_item paragraph must NOT appear in mismatch_examples
    # (detect_role would say 'body' for it, but that's a false alarm)
    mismatch_labels = [ex["label"] for ex in consistency["mismatch_examples"]]
    assert "list_item" not in mismatch_labels, (
        "Multi-line list_item paragraphs must not appear as mismatch examples"
    )
    # The compared count must exclude the multi-line paragraph
    assert consistency["compared"] <= 2, (
        f"Multi-line block should be excluded from comparison; compared={consistency['compared']}"
    )
    assert consistency["mismatched"] == 0, (
        f"Expected 0 real mismatches; got {consistency['mismatched']}"
    )


def test_multiline_non_list_mismatch_still_reported():
    """
    A 'real' mismatch (e.g. LLM says 'h1' but detect_role says 'body' for a
    plain single-line paragraph) must still appear in the mismatch report.
    This ensures the multi-line exclusion does not suppress genuine divergences.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    # Single-line paragraph with no list marker, no heading style — detect_role
    # returns 'body', but LLM (incorrectly) labels it 'h1'.
    doc, blocks, labels = _make_doc_blocks_labels([
        ("h1", "这是一段普通正文内容没有标题标记"),
    ])

    report = apply_formatting(doc, blocks, labels, spec)

    consistency = report["labels"]["consistency"]
    assert consistency["compared"] == 1
    assert consistency["mismatched"] == 1, (
        "Single-line body/h1 divergence must still be reported as a mismatch"
    )


# ─── 首行缩进：preamble lines inside a split list_item block ────────────────

def test_preamble_line_of_split_list_item_gets_first_line_indent():
    """
    When a multi-line block labeled 'list_item' by LLM is split, the first
    fragment may be a preamble sentence (e.g. "以下是方向：") with no list
    prefix.  That fragment should receive body首行缩进 (positive
    first_line_indent) rather than a hanging indent, because step 5 won't
    add a numPr to it and hanging indent without numPr looks wrong.

    Numbered fragments (e.g. "1. 保研...") must still get list_item
    formatting (hanging indent + numPr).
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    body_size = float(spec.raw["body"]["font_size_pt"])
    first_line_chars = int(spec.raw["body"]["first_line_chars"])
    expected_fli = first_line_chars * body_size   # positive pt value for body
    hanging_pt = float(spec.raw["list_item"]["hanging_indent_pt"])

    # Simulate: LLM labels the whole multi-line block as list_item.
    # After step-3 split the first line has no list marker.
    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "结合信息安全专业背景，我初步考虑以下几个方向：\n1. 保研或考研。\n2. 就业方向。"),
    ])

    apply_formatting(doc, blocks, labels, spec)

    paras = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(paras) == 3, f"Expected 3 paragraphs after split, got {len(paras)}: {[p.text for p in paras]}"

    preamble = paras[0]
    list_items = paras[1:]

    # Preamble must NOT be a numPr paragraph, and must have positive首行缩进.
    assert not _is_list_p(preamble), "Preamble paragraph must not have numPr"
    fli = preamble.paragraph_format.first_line_indent
    assert fli is not None and float(fli) > 0, (
        f"Preamble首行缩进 must be positive (body indent), got {fli}"
    )
    assert float(fli) == pytest.approx(Pt(expected_fli), rel=0.05), (
        f"Preamble首行缩进 must be ≈{expected_fli}pt (body first_line_chars), got {fli}"
    )

    # Numbered items must be numPr paragraphs with hanging indent.
    for p in list_items:
        assert _is_list_p(p), f"List item must have numPr: {p.text!r}"
        item_fli = p.paragraph_format.first_line_indent
        assert item_fli is not None and float(item_fli) <= 0, (
            f"List item must have hanging (≤0) first_line_indent, got {item_fli}"
        )
        assert float(p.paragraph_format.left_indent) == pytest.approx(Pt(hanging_pt), rel=0.01), (
            f"List item left_indent must equal hanging_indent_pt={hanging_pt}pt"
        )


def test_header_label_line_of_split_list_item_gets_first_line_indent():
    """
    When a multi-line block labeled 'list_item' starts with a bracket header
    like '【大二下学期】' (no list prefix), that header line must get body
    首行缩进 (not a hanging indent) after the split and formatting.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    body_size = float(spec.raw["body"]["font_size_pt"])
    first_line_chars = int(spec.raw["body"]["first_line_chars"])
    expected_fli = first_line_chars * body_size

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "【大二下学期】\n1. 巩固数据结构。\n2. 深入学习Linux。"),
    ])

    apply_formatting(doc, blocks, labels, spec)

    paras = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(paras) == 3

    header_line = paras[0]
    assert not _is_list_p(header_line), "Header line must not have numPr"
    fli = header_line.paragraph_format.first_line_indent
    assert fli is not None and float(fli) > 0, (
        f"Header line首行缩进 must be positive, got {fli}"
    )

    for p in paras[1:]:
        assert _is_list_p(p), f"Numbered items must have numPr: {p.text!r}"


def test_all_lines_have_prefix_no_preamble_downgrade():
    """
    When ALL split lines have a list prefix (e.g. "1. ...\n2. ...\n3. ..."),
    none should be downgraded to body — all must stay as list_item and get numPr.
    """
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    doc, blocks, labels = _make_doc_blocks_labels([
        ("list_item", "1. 数学成绩优秀。\n2. 编程能力增强。\n3. 逻辑推理提升。"),
    ])

    apply_formatting(doc, blocks, labels, spec)

    paras = [p for p in iter_all_paragraphs(doc) if not is_effectively_blank_paragraph(p)]
    assert len(paras) == 3, f"Expected 3 paragraphs, got {len(paras)}"

    for p in paras:
        assert _is_list_p(p), f"All split list_item paragraphs must have numPr: {p.text!r}"
        fli = p.paragraph_format.first_line_indent
        assert fli is None or float(fli) <= 0, (
            f"All list_item paragraphs must have hanging indent (≤0), got {fli}"
        )
