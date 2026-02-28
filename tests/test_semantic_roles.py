# tests/test_semantic_roles.py
"""
Tests for semantic role preservation (abstract/keyword/reference/footer/list_item)
in formatter, spec loader, and all specialized template specs.
"""
from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.spec import load_spec
from core.formatter import apply_formatting, detect_role


SPECS_DIR = Path(__file__).resolve().parents[1] / "specs"


def _make_doc_with_roles(role_texts: list) -> tuple:
    """Create a Document with one paragraph per (role, text) pair, return (doc, blocks, labels)."""
    from core.docx_utils import iter_all_paragraphs
    from core.parser import Block

    doc = Document()
    for _, text in role_texts:
        doc.add_paragraph(text)

    blocks = [
        Block(block_id=i + 1, kind="paragraph", text=p.text, paragraph_index=i)
        for i, p in enumerate(iter_all_paragraphs(doc))
    ]
    labels = {b.block_id: role for b, (role, _) in zip(blocks, role_texts)}
    labels["_source"] = "test"
    return doc, blocks, labels


def _mark_paragraph_as_numbered_list(paragraph) -> None:
    """Mark a paragraph as a numbered list item via numPr XML for detect_role tests."""
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = 1


def test_semantic_roles_formatted_not_as_unknown():
    """abstract/keyword/reference/footer/list_item must appear in formatted counts, not unknown_as_body."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))

    semantic_roles = ["abstract", "keyword", "reference", "footer", "list_item"]
    role_texts = [(role, f"这是{role}段落内容示例。") for role in semantic_roles]

    doc, blocks, labels = _make_doc_with_roles(role_texts)
    report = apply_formatting(doc, blocks, labels, spec)

    counts = report["formatted"]["counts"]
    for role in semantic_roles:
        assert counts.get(role, 0) >= 1, (
            f"Role '{role}' should appear in formatted counts but got: {counts}"
        )
    assert counts.get("unknown_as_body", 0) == 0, (
        f"No paragraph should fall through to unknown_as_body: {counts}"
    )


def test_all_spec_files_load_successfully():
    """All YAML spec files in specs/ should load without error."""
    for yaml_path in SPECS_DIR.glob("*.yaml"):
        spec = load_spec(str(yaml_path))
        assert spec is not None, f"Failed to load {yaml_path.name}"
        # Must have required top-level keys
        assert "body" in spec.raw
        assert "heading" in spec.raw
        # Must have semantic role defaults filled in
        for role in ("abstract", "keyword", "reference", "footer"):
            assert role in spec.raw, f"'{role}' key missing in {yaml_path.name}"
            assert "font_size_pt" in spec.raw[role], (
                f"font_size_pt missing in {role} section of {yaml_path.name}"
            )


def test_reference_has_hanging_indent_in_academic_spec():
    """Academic spec should configure reference with non-zero hanging indent."""
    spec = load_spec(str(SPECS_DIR / "academic.yaml"))
    ref = spec.raw["reference"]
    assert float(ref.get("hanging_indent_pt", 0)) > 0, (
        "Academic reference should have hanging_indent_pt > 0 (GB/T 7714 style)"
    )


def test_abstract_italic_in_default_spec():
    """Default spec abstract should be italic to visually distinguish it from body."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    assert spec.raw["abstract"]["italic"] is True


def test_detect_role_semantic_patterns():
    """detect_role should recognize semantic text patterns and numbered list paragraphs."""
    doc = Document()
    p_abs = doc.add_paragraph("摘要：这是摘要内容。")
    p_kw = doc.add_paragraph("关键词：测试；排版")
    p_ref = doc.add_paragraph("参考文献")
    p_list = doc.add_paragraph("第一条")

    _mark_paragraph_as_numbered_list(p_list)

    assert detect_role(p_abs) == "abstract"
    assert detect_role(p_kw) == "keyword"
    assert detect_role(p_ref) == "reference"
    assert detect_role(p_list) == "list_item"


def test_detect_role_semantic_patterns_edge_variants():
    """detect_role should support case and punctuation variants for semantic patterns."""
    doc = Document()
    assert detect_role(doc.add_paragraph("ABSTRACT This is abstract content.")) == "abstract"
    assert detect_role(doc.add_paragraph("Keywords test, parser")) == "keyword"
    assert detect_role(doc.add_paragraph("  references  ")) == "reference"


def test_unknown_label_falls_back_to_semantic_detect_role():
    """unknown labels should fall back to semantic detect_role instead of unknown_as_body."""
    spec = load_spec(str(SPECS_DIR / "default.yaml"))
    role_texts = [("unknown", "摘要：这里是摘要内容。")]
    doc, blocks, labels = _make_doc_with_roles(role_texts)

    report = apply_formatting(doc, blocks, labels, spec)
    counts = report["formatted"]["counts"]

    assert counts.get("abstract", 0) == 1
    assert counts.get("unknown_as_body", 0) == 0
