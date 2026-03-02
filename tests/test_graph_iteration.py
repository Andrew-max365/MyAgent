# tests/test_graph_iteration.py
"""Tests for graph iteration stopping conditions using mocked LangGraph nodes."""
from __future__ import annotations

from unittest.mock import MagicMock, patch

import pytest

from agent.graph.nodes import retry_router
from agent.graph.react_schemas import GraphState


def _make_state(**overrides) -> GraphState:
    base: GraphState = {
        "input_path": "in.docx",
        "output_path": "out.docx",
        "spec_path": "specs/default.yaml",
        "label_mode": "rule",
        "max_iters": 3,
        "current_iter": 0,
        "thoughts": [],
        "actions": [],
        "observations": [],
        "errors": [],
        "passed": False,
        "finished": False,
        "report": {},
        "blocks": None,
        "labels": None,
        "doc": None,
    }
    base.update(overrides)
    return base


# --- retry_router tests ---

def test_retry_router_end_when_passed():
    state = _make_state(passed=True, current_iter=1)
    assert retry_router(state) == "end"


def test_retry_router_end_when_max_iters_reached():
    state = _make_state(passed=False, current_iter=3, max_iters=3)
    assert retry_router(state) == "end"


def test_retry_router_reason_when_not_passed():
    state = _make_state(passed=False, current_iter=1, max_iters=3)
    assert retry_router(state) == "reason"


def test_retry_router_end_when_iter_exceeds_max():
    state = _make_state(passed=False, current_iter=5, max_iters=3)
    assert retry_router(state) == "end"


# --- reason_node tests ---

def test_reason_node_increments_iter():
    from agent.graph.nodes import reason_node
    state = _make_state(current_iter=0, errors=[])
    result = reason_node(state)
    assert result["current_iter"] == 1
    assert len(result["thoughts"]) == 1
    assert len(result["actions"]) == 1


def test_reason_node_appends_to_existing():
    from agent.graph.nodes import reason_node
    state = _make_state(
        current_iter=1,
        thoughts=["first thought"],
        actions=[[{"action_type": "no_op", "block_id": -1, "params": {}}]],
        errors=["some error"],
    )
    result = reason_node(state)
    assert result["current_iter"] == 2
    assert len(result["thoughts"]) == 2
    assert "some error" in result["thoughts"][1]


def test_reason_node_produces_no_op_action():
    from agent.graph.nodes import reason_node
    state = _make_state(current_iter=0)
    result = reason_node(state)
    actions = result["actions"][-1]
    assert len(actions) == 1
    assert actions[0]["action_type"] == "no_op"


# --- act_node tests ---

def test_act_node_applies_set_role(tmp_path):
    """act_node should override label when action_type=set_role."""
    from agent.graph.nodes import act_node
    from docx import Document
    from core.parser import parse_docx_to_blocks

    doc = Document()
    doc.add_paragraph("Title paragraph")
    doc.add_paragraph("Body paragraph")
    docx_path = str(tmp_path / "test.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)

    state = _make_state(
        doc=doc2,
        blocks=blocks,
        actions=[[{"action_type": "set_role", "block_id": 0, "params": {"role": "h1"}}]],
    )
    result = act_node(state)
    assert result["labels"][0] == "h1"


def test_act_node_applies_fix_heading_level(tmp_path):
    from agent.graph.nodes import act_node
    from docx import Document
    from core.parser import parse_docx_to_blocks

    doc = Document()
    doc.add_paragraph("Section heading")
    docx_path = str(tmp_path / "test2.docx")
    doc.save(docx_path)
    doc2, blocks = parse_docx_to_blocks(docx_path)

    state = _make_state(
        doc=doc2,
        blocks=blocks,
        actions=[[{"action_type": "fix_heading_level", "block_id": 0, "params": {"level": 2}}]],
    )
    result = act_node(state)
    assert result["labels"][0] == "h2"
