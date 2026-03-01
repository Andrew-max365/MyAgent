from pathlib import Path
import sys

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.formatter import _split_body_paragraphs_on_linebreaks


def test_split_linebreaks_in_table_cell_body_paragraph():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.cell(0, 0).paragraphs[0]
    p.text = "第一行\n第二行"

    created = _split_body_paragraphs_on_linebreaks(doc, role_getter=lambda _: "body")

    paras = table.cell(0, 0).paragraphs
    assert created == 1
    assert [para.text for para in paras] == ["第一行", "第二行"]


def test_split_linebreaks_keeps_non_body_unchanged():
    doc = Document()
    p = doc.add_paragraph("标题\n副标题")

    created = _split_body_paragraphs_on_linebreaks(doc, role_getter=lambda _: "h1")

    assert created == 0
    assert p.text == "标题\n副标题"


def test_split_linebreaks_in_table_cell_list_item_paragraph():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.cell(0, 0).paragraphs[0]
    p.text = "1）第一行\n2）第二行\n3）第三行"

    created = _split_body_paragraphs_on_linebreaks(doc, role_getter=lambda _: "list_item")

    paras = table.cell(0, 0).paragraphs
    assert created == 2
    assert [para.text for para in paras] == ["1）第一行", "2）第二行", "3）第三行"]


def test_split_linebreaks_in_table_cell_list_item_carriage_return():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.cell(0, 0).paragraphs[0]
    p.text = "1）第一行\r2）第二行\r3）第三行"

    created = _split_body_paragraphs_on_linebreaks(doc, role_getter=lambda _: "list_item")

    paras = table.cell(0, 0).paragraphs
    assert created == 2
    assert [para.text for para in paras] == ["1）第一行", "2）第二行", "3）第三行"]
