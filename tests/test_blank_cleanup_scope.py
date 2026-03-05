from pathlib import Path
import sys

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))

from core.formatter import _cleanup_consecutive_blanks, _delete_blanks_after_roles


def test_cleanup_consecutive_blanks_handles_table_cells_independently():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)

    c1 = table.cell(0, 0)
    c1.paragraphs[0].text = ""
    c1.add_paragraph("")
    c1.add_paragraph("")

    c2 = table.cell(0, 1)
    c2.paragraphs[0].text = ""
    c2.add_paragraph("内容")

    deleted = _cleanup_consecutive_blanks(doc, max_keep=1)

    assert deleted == 2
    assert len(c1.paragraphs) == 1
    assert [p.text for p in c2.paragraphs] == ["", "内容"]


def test_delete_blanks_after_roles_within_same_container_only():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p_title = cell.paragraphs[0]
    p_title.text = "标题"
    p_blank = cell.add_paragraph("")
    p_body = cell.add_paragraph("正文")

    deleted = _delete_blanks_after_roles(
        doc,
        roles={"h1"},
        role_getter=lambda p: "h1" if p.text == "标题" else "body",
    )

    assert deleted == 1
    assert [p.text for p in cell.paragraphs] == ["标题", "正文"]
