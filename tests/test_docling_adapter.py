# tests/test_docling_adapter.py
"""Tests for Docling adapter fallback behavior."""
from __future__ import annotations

import warnings
from unittest.mock import MagicMock, patch

import pytest

from core.docling_adapter import (
    DOCLING_AVAILABLE,
    DOCLING_LABEL_MAP,
    _map_docling_to_blocks,
    parse_with_fallback,
)


def test_docling_label_map_has_expected_keys():
    assert "title" in DOCLING_LABEL_MAP
    assert "section_header" in DOCLING_LABEL_MAP
    assert "list_item" in DOCLING_LABEL_MAP
    assert DOCLING_LABEL_MAP["title"] == "h1"
    assert DOCLING_LABEL_MAP["section_header"] == "h2"
    # New labels from DocItemLabel enum
    assert "paragraph" in DOCLING_LABEL_MAP, "DocItemLabel.PARAGRAPH must be mapped"
    assert DOCLING_LABEL_MAP["paragraph"] == "body"
    assert "code" in DOCLING_LABEL_MAP
    assert "formula" in DOCLING_LABEL_MAP


def test_map_docling_to_blocks_empty():
    mock_doc = MagicMock()
    mock_doc.texts = []
    result = _map_docling_to_blocks(mock_doc)
    assert result == []


def test_map_docling_to_blocks_maps_labels():
    item = MagicMock()
    item.label = "title"
    item.text = "Hello"
    mock_doc = MagicMock()
    mock_doc.texts = [item]
    result = _map_docling_to_blocks(mock_doc)
    assert len(result) == 1
    assert result[0]["role"] == "h1"
    assert result[0]["text"] == "Hello"


def test_map_docling_to_blocks_unknown_label():
    item = MagicMock()
    item.label = "unknown_type"
    item.text = "some text"
    mock_doc = MagicMock()
    mock_doc.texts = [item]
    result = _map_docling_to_blocks(mock_doc)
    assert result[0]["role"] == "body"


def test_map_docling_to_blocks_handles_exception():
    mock_doc = MagicMock()
    mock_doc.texts = None  # will cause TypeError when iterating
    # Should not raise
    result = _map_docling_to_blocks(mock_doc)
    assert isinstance(result, list)


def test_map_docling_to_blocks_paragraph_label():
    """DocItemLabel.PARAGRAPH (used for Word processor paragraphs) should map to body."""
    item = MagicMock()
    item.label = "paragraph"
    item.text = "Some body text"
    mock_doc = MagicMock()
    mock_doc.texts = [item]
    result = _map_docling_to_blocks(mock_doc)
    assert result[0]["role"] == "body"


def test_map_docling_to_blocks_section_header_level_aware():
    """SectionHeaderItem with level attribute should produce h1/h2/h3."""
    items = []
    for lvl, expected_role in [(1, "h1"), (2, "h2"), (3, "h3"), (4, "h3")]:
        item = MagicMock()
        item.label = "section_header"
        item.level = lvl
        item.text = f"Header level {lvl}"
        items.append((item, expected_role))

    mock_doc = MagicMock()
    mock_doc.texts = [i for i, _ in items]
    result = _map_docling_to_blocks(mock_doc)
    for (item, expected), got in zip(items, result):
        assert got["role"] == expected, f"level={item.level}: expected {expected}, got {got['role']}"


def test_map_docling_to_blocks_section_header_no_level_defaults_h2():
    """SectionHeaderItem without level attribute defaults to h2."""
    item = MagicMock(spec=["label", "text"])  # no 'level' attr
    item.label = "section_header"
    item.text = "Header"
    mock_doc = MagicMock()
    mock_doc.texts = [item]
    result = _map_docling_to_blocks(mock_doc)
    assert result[0]["role"] == "h2"


def test_parse_with_fallback_uses_original_parser_when_docling_disabled(tmp_path):
    from docx import Document
    docx_path = str(tmp_path / "sample.docx")
    doc = Document()
    doc.add_paragraph("Test paragraph")
    doc.save(docx_path)

    doc_result, blocks = parse_with_fallback(docx_path, use_docling=False)
    assert doc_result is not None
    assert isinstance(blocks, list)


def test_parse_with_fallback_falls_back_on_docling_failure(tmp_path):
    from docx import Document
    docx_path = str(tmp_path / "sample.docx")
    doc = Document()
    doc.add_paragraph("Test paragraph")
    doc.save(docx_path)

    # Even if use_docling=True but docling not installed, fallback happens silently
    doc_result, blocks = parse_with_fallback(docx_path, use_docling=True)
    assert doc_result is not None
    assert isinstance(blocks, list)


def test_parse_with_fallback_warns_on_docling_exception(tmp_path):
    """When Docling raises, a warning is emitted and parser fallback occurs."""
    from docx import Document
    docx_path = str(tmp_path / "sample.docx")
    doc = Document()
    doc.add_paragraph("Test paragraph")
    doc.save(docx_path)

    with patch("core.docling_adapter.DOCLING_AVAILABLE", True), \
         patch("core.docling_adapter.parse_with_docling", side_effect=RuntimeError("mock fail")):
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            doc_result, blocks = parse_with_fallback(docx_path, use_docling=True)
            assert any("Docling" in str(warning.message) for warning in w)

    assert doc_result is not None
