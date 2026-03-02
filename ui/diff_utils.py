# ui/diff_utils.py
"""
Utilities for generating diff views and applying LLM proofread suggestions.

Provides:
  - DiffItem: a numbered, human-readable description of a single text change
  - build_diff_items(): convert raw proofread issue dicts → [DiffItem]
  - parse_rejected_numbers(): parse user text like "不要修改#3 #5" → {3, 5}
  - apply_proofread_issues(): apply text replacements to a python-docx Document
  - apply_and_save_proofread(): convenience wrapper – load bytes, apply, save bytes
  - generate_structural_diff(): markdown summary of structural formatting changes
"""
from __future__ import annotations

import io as _io
import os
import re
import tempfile
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Set, Tuple


# ---------------------------------------------------------------------------
# DiffItem – one numbered proofread change
# ---------------------------------------------------------------------------

_ISSUE_ICONS: Dict[str, str] = {
    "typo": "🔤",
    "punctuation": "🔡",
    "standardization": "📏",
}
_SEVERITY_ICONS: Dict[str, str] = {
    "low": "🟢",
    "medium": "🟡",
    "high": "🔴",
}
_ISSUE_LABELS: Dict[str, str] = {
    "typo": "错别字",
    "punctuation": "标点符号",
    "standardization": "规范性",
}


@dataclass
class DiffItem:
    """A numbered diff item representing a single proofread text change."""

    number: int
    issue_type: str          # "typo" | "punctuation" | "standardization"
    severity: str            # "low" | "medium" | "high"
    para_idx: Optional[int]  # 0-based paragraph index, if known
    evidence: str            # original text fragment
    suggestion: str          # suggested replacement
    rationale: str           # explanation

    def to_markdown(self) -> str:
        """Render as a compact markdown diff card."""
        icon = _ISSUE_ICONS.get(self.issue_type, "📝")
        sev = _SEVERITY_ICONS.get(self.severity, "")
        label = _ISSUE_LABELS.get(self.issue_type, self.issue_type)
        location = f" (段落 {self.para_idx})" if self.para_idx is not None else ""
        return (
            f"**#{self.number}** {icon}{sev} `{label}`{location}\n"
            f"  - ❌ 原文片段：`{self.evidence}`\n"
            f"  - ✅ 建议修改：`{self.suggestion}`\n"
            f"  - 💡 说明：{self.rationale}"
        )


def build_diff_items(issues: List[dict]) -> List[DiffItem]:
    """
    Convert a list of raw proofread issue dicts (from report['llm_proofread']['issues'])
    into numbered DiffItem objects.
    """
    items: List[DiffItem] = []
    for i, issue in enumerate(issues, start=1):
        evidence = (issue.get("evidence") or "").strip()
        suggestion = (issue.get("suggestion") or "").strip()
        if not evidence or not suggestion:
            continue
        items.append(
            DiffItem(
                number=i,
                issue_type=issue.get("issue_type", "standardization"),
                severity=issue.get("severity", "low"),
                para_idx=issue.get("paragraph_index"),
                evidence=evidence,
                suggestion=suggestion,
                rationale=(issue.get("rationale") or ""),
            )
        )
    return items


# ---------------------------------------------------------------------------
# Parsing user rejection commands
# ---------------------------------------------------------------------------

# Patterns that signal "accept all" (i.e., no rejections)
_ACCEPT_ALL_PATTERNS = re.compile(
    r"(全部接受|全部同意|全部确认|确认所有|accept\s*all|接受全部|同意全部|应用全部|全部应用|^确认$|^同意$|^好的?$|^ok$)",
    re.IGNORECASE,
)
# Patterns that signal "reject all"
_REJECT_ALL_PATTERNS = re.compile(
    r"(全部不要|全部拒绝|不要任何|全部保留原文|reject\s*all|不接受任何|全部回退)",
    re.IGNORECASE,
)


def parse_rejected_numbers(user_text: str, total: int) -> Tuple[Set[int], str]:
    """
    Parse user text to find which change numbers are being rejected.

    Returns (rejected_set, intent) where intent is one of:
      "accept_all"  – apply every change
      "reject_all"  – apply no change
      "partial"     – apply some, reject the identified numbers

    Examples handled:
      "不要修改#3"            → ({3}, "partial")
      "不要改#3 和 #5"        → ({3, 5}, "partial")
      "reject 3, 5"           → ({3, 5}, "partial")
      "不要第3条和第5条"      → ({3, 5}, "partial")
      "全部接受"              → ({}, "accept_all")
      "全部不要"              → (set(range(1, total+1)), "reject_all")
    """
    if _ACCEPT_ALL_PATTERNS.search(user_text):
        return set(), "accept_all"
    if _REJECT_ALL_PATTERNS.search(user_text):
        return set(range(1, total + 1)), "reject_all"

    rejected: Set[int] = set()

    # Match explicit #N patterns first (highest confidence)
    for m in re.finditer(r"#\s*(\d+)", user_text):
        n = int(m.group(1))
        if 1 <= n <= total:
            rejected.add(n)

    # Match "第N条" / "第N个" patterns
    for m in re.finditer(r"第\s*(\d+)\s*[条个项]", user_text):
        n = int(m.group(1))
        if 1 <= n <= total:
            rejected.add(n)

    # If nothing matched yet, try standalone numbers (less confident – only when
    # "rejection" context is present in the message)
    if not rejected:
        rejection_keywords = re.compile(
            r"(不要|拒绝|取消|不接受|skip|reject|ignore|不改|保留原文|撤销)",
            re.IGNORECASE,
        )
        if rejection_keywords.search(user_text):
            for m in re.finditer(r"\b(\d+)\b", user_text):
                n = int(m.group(1))
                if 1 <= n <= total:
                    rejected.add(n)

    intent = "partial" if rejected else "accept_all"
    return rejected, intent


# ---------------------------------------------------------------------------
# Applying proofread issues to a python-docx Document
# ---------------------------------------------------------------------------

def _replace_text_in_paragraph(para, old_text: str, new_text: str) -> bool:
    """
    Replace the first occurrence of *old_text* with *new_text* inside *para*.

    Tries single-run replacement first; falls back to a multi-run approach where
    all run texts are merged, the substitution is performed, and the combined
    text is placed back in the first run (remaining runs are cleared).

    Returns True if a replacement was made.
    """
    # Fast path: the text is entirely within one run
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True

    # Slow path: text spans multiple runs
    runs = list(para.runs)
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    pos = full.find(old_text)
    if pos < 0:
        return False
    new_full = full[:pos] + new_text + full[pos + len(old_text):]
    runs[0].text = new_full
    for run in runs[1:]:
        run.text = ""
    return True


def apply_proofread_issues(
    doc,
    issues: List[dict],
    excluded_numbers: Optional[Set[int]] = None,
) -> int:
    """
    Apply proofread issues (text replacements) to a python-docx Document in-place.

    :param doc:               python-docx Document object (modified in-place)
    :param issues:            list of raw issue dicts (from report['llm_proofread']['issues'])
    :param excluded_numbers:  1-based set of issue numbers to skip
    :return:                  number of replacements successfully applied
    """
    if not issues:
        return 0

    from core.docx_utils import iter_all_paragraphs

    excluded = excluded_numbers or set()
    paragraphs = iter_all_paragraphs(doc)
    applied = 0

    for i, issue in enumerate(issues, start=1):
        if i in excluded:
            continue
        evidence = (issue.get("evidence") or "").strip()
        suggestion = (issue.get("suggestion") or "").strip()
        if not evidence or not suggestion or evidence == suggestion:
            continue

        para_idx = issue.get("paragraph_index")
        if para_idx is not None and 0 <= para_idx < len(paragraphs):
            if _replace_text_in_paragraph(paragraphs[para_idx], evidence, suggestion):
                applied += 1
                continue
        # Fallback: search all paragraphs
        for para in paragraphs:
            if evidence in para.text:
                if _replace_text_in_paragraph(para, evidence, suggestion):
                    applied += 1
                    break

    return applied


def apply_and_save_proofread(
    output_bytes: bytes,
    issues: List[dict],
    excluded_numbers: Optional[Set[int]] = None,
) -> Tuple[bytes, int]:
    """
    Load a docx from *output_bytes*, apply proofread issues (excluding
    *excluded_numbers*), and return the modified docx bytes + count applied.
    """
    from docx import Document
    from core.writer import save_docx

    tmp_in = tmp_out = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(output_bytes)
            tmp_in = f.name
        tmp_out = tmp_in + "_proofread.docx"

        doc = Document(tmp_in)
        applied = apply_proofread_issues(doc, issues, excluded_numbers)
        save_docx(doc, tmp_out)

        with open(tmp_out, "rb") as f:
            result_bytes = f.read()

        return result_bytes, applied

    finally:
        for p in (tmp_in, tmp_out):
            if p:
                try:
                    os.remove(p)
                except OSError:
                    pass


# ---------------------------------------------------------------------------
# Markdown side-by-side diff table
# ---------------------------------------------------------------------------


def generate_diff_markdown(diff_items: List[DiffItem]) -> str:
    """
    Generate a Markdown string containing a side-by-side diff table for *diff_items*.

    The table has columns: #, 类型, ❌ 原文, ✅ 建议, 说明.
    Rendered directly in Chainlit message content as a GFM table.
    """
    lines: List[str] = [
        "| # | 类型 | ❌ 原文 | ✅ 建议 | 说明 |",
        "|---|------|--------|--------|------|",
    ]
    for item in diff_items:
        icon = _ISSUE_ICONS.get(item.issue_type, "📝")
        sev = _SEVERITY_ICONS.get(item.severity, "")
        label = _ISSUE_LABELS.get(item.issue_type, item.issue_type)
        location = f" (段落{item.para_idx})" if item.para_idx is not None else ""

        # Escape pipe chars in cell content so they don't break the table
        def _esc(s: str) -> str:
            return s.replace("|", "\\|").replace("\n", " ")

        lines.append(
            f"| **#{item.number}** "
            f"| {icon}{sev} {_esc(label)}{_esc(location)} "
            f"| `{_esc(item.evidence)}` "
            f"| `{_esc(item.suggestion)}` "
            f"| {_esc(item.rationale)} |"
        )
    return "\n".join(lines)


# Keep the old HTML generator for backwards compatibility if needed elsewhere
def generate_diff_html(diff_items: List[DiffItem]) -> str:
    """Alias that returns markdown; HTML rendering was replaced with markdown tables."""
    return generate_diff_markdown(diff_items)


def generate_diff_cards_markdown(
    diff_items: List[DiffItem],
    paragraph_texts: Optional[List[str]] = None,
) -> str:
    """
    Generate a GFM markdown string with one "card" per diff item that renders
    visually in the Chainlit browser UI without any download:

    - Paragraph context shown as a Markdown blockquote.
    - Original fragment rendered with GFM **~~strikethrough~~** (displays as
      red-ish struck-out text in most renderers).
    - Suggested replacement rendered as **bold**.
    - Rationale shown as italic below the diff line.

    :param diff_items:       List of DiffItem objects to render.
    :param paragraph_texts:  Optional list of raw paragraph strings from the
                             formatted document (indexed by paragraph position).
                             When supplied, the relevant paragraph is quoted above
                             the diff line so the reviewer sees the full context.
    """
    blocks: List[str] = []
    for item in diff_items:
        icon = _ISSUE_ICONS.get(item.issue_type, "📝")
        sev = _SEVERITY_ICONS.get(item.severity, "")
        label = _ISSUE_LABELS.get(item.issue_type, item.issue_type)
        loc = f"（段落 {item.para_idx}）" if item.para_idx is not None else ""

        # ── Escape markdown special chars inside variable content ────────────
        def _esc(s: str) -> str:
            return (
                s.replace("\\", "\\\\")
                .replace("`", "\\`")
                .replace("*", "\\*")
                .replace("_", "\\_")
                .replace("~", "\\~")
                .replace("[", "\\[")
                .replace("]", "\\]")
                .replace("\n", " ")
            )

        card_lines: List[str] = [
            f"**#{item.number}** {icon}{sev} `{label}`{loc}",
        ]

        # Paragraph context as a blockquote
        para_idx = item.para_idx
        if (
            para_idx is not None
            and paragraph_texts is not None
            and 0 <= para_idx < len(paragraph_texts)
        ):
            ctx = paragraph_texts[para_idx].strip()
            if len(ctx) > 120:
                ctx = ctx[:120] + "…"
            if ctx:
                card_lines.append(f"> {ctx}")

        # Diff line: ~~evidence~~ → **suggestion**
        card_lines.append(
            f"~~{_esc(item.evidence)}~~ → **{_esc(item.suggestion)}**"
        )

        # Rationale
        if item.rationale:
            card_lines.append(f"💡 _{_esc(item.rationale)}_")

        blocks.append("\n\n".join(card_lines))

    return "\n\n---\n\n".join(blocks)


# ---------------------------------------------------------------------------
# Redline Word document (tracked-changes-style preview)
# ---------------------------------------------------------------------------

_DEL_COLOR = (0xC0, 0x00, 0x00)   # red
_INS_COLOR = (0x00, 0x70, 0x00)   # green
_NOTE_COLOR = (0x66, 0x66, 0x66)  # grey
_ISSUE_LABELS_ZH: Dict[str, str] = {
    "typo": "错别字",
    "punctuation": "标点符号",
    "standardization": "规范性",
}
_SEVERITY_ZH: Dict[str, str] = {
    "low": "低",
    "medium": "中",
    "high": "高",
}


def generate_redline_docx(diff_items: List[DiffItem], out_bytes: bytes) -> bytes:
    """
    Generate a "redline" Word document that shows each LLM proofread suggestion
    as a visual before/after diff:

    - Original text: **red strikethrough**
    - Suggested replacement: **green underline**
    - Paragraph context from *out_bytes* (the formatted document) is shown above
      each change so the reviewer knows exactly where it falls.

    Returns the redline .docx as raw bytes ready to attach as a cl.File.
    """
    from docx import Document
    from docx.shared import RGBColor, Pt

    # ── Load paragraph context from formatted document ───────────────────────
    paragraphs: list = []
    if out_bytes:
        _tmp = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                f.write(out_bytes)
                _tmp = f.name
            from core.docx_utils import iter_all_paragraphs
            paragraphs = iter_all_paragraphs(Document(_tmp))
        except Exception:
            pass
        finally:
            if _tmp:
                try:
                    os.remove(_tmp)
                except OSError:
                    pass

    # ── Build the redline document ───────────────────────────────────────────
    doc = Document()
    doc.add_heading("LLM 校对建议 — 修订预览", 1)

    for item in diff_items:
        label = _ISSUE_LABELS_ZH.get(item.issue_type, item.issue_type)
        sev = _SEVERITY_ZH.get(item.severity, item.severity)
        loc = f"（段落 {item.para_idx}）" if item.para_idx is not None else ""

        # ── Section heading ──────────────────────────────────────────────────
        p_head = doc.add_paragraph()
        r_head = p_head.add_run(f"#{item.number}  [{label} · 严重性: {sev}]{loc}")
        r_head.bold = True

        # ── Paragraph context ────────────────────────────────────────────────
        para_idx = item.para_idx
        if para_idx is not None and 0 <= para_idx < len(paragraphs):
            ctx = paragraphs[para_idx].text
            if len(ctx) > 150:
                ctx = ctx[:150] + "…"
            p_ctx = doc.add_paragraph()
            r_ctx = p_ctx.add_run("原段落：" + ctx)
            r_ctx.font.color.rgb = RGBColor(*_NOTE_COLOR)
            r_ctx.font.size = Pt(9)

        # ── Diff line: red strikethrough → green underline ───────────────────
        p_diff = doc.add_paragraph()
        p_diff.add_run("修改：")
        r_del = p_diff.add_run(item.evidence)
        r_del.font.color.rgb = RGBColor(*_DEL_COLOR)
        r_del.font.strike = True
        p_diff.add_run("  →  ")
        r_ins = p_diff.add_run(item.suggestion)
        r_ins.font.color.rgb = RGBColor(*_INS_COLOR)
        r_ins.font.underline = True

        # ── Rationale ────────────────────────────────────────────────────────
        if item.rationale:
            p_note = doc.add_paragraph()
            r_note = p_note.add_run(f"说明：{item.rationale}")
            r_note.font.color.rgb = RGBColor(*_NOTE_COLOR)
            r_note.font.size = Pt(10)

        doc.add_paragraph()  # blank spacer between items

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Structural diff summary
# ---------------------------------------------------------------------------

def generate_structural_diff(report: dict) -> str:
    """
    Generate a markdown summary of the structural formatting changes recorded
    in the *actions* section of the format report.
    """
    actions = report.get("actions", {})
    if not actions:
        return ""

    lines: List[str] = []

    _MAPPING = [
        ("h1_applied",         "一级标题样式应用"),
        ("h2_applied",         "二级标题样式应用"),
        ("h3_applied",         "三级标题样式应用"),
        ("body_applied",       "正文样式应用"),
        ("caption_applied",    "题注样式应用"),
        ("abstract_applied",   "摘要样式应用"),
        ("keyword_applied",    "关键词样式应用"),
        ("reference_applied",  "参考文献样式应用"),
        ("list_converted",     "列表项转换（numPr）"),
        ("tables_autofitted",  "表格自动适配"),
        ("blank_removed",      "空段落清理"),
    ]

    for key, label in _MAPPING:
        n = actions.get(key, 0)
        if n:
            lines.append(f"- **{n}** 处 {label}")

    return "\n".join(lines)
